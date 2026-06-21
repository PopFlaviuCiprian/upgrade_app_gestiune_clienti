import sys
import tkinter as tk
from tkinter import ttk, messagebox, simpledialog
import requests
import webbrowser
from tkcalendar import DateEntry
from datetime import datetime, date
from datetime import datetime, timedelta
from dateutil.relativedelta import relativedelta
import csv
from tkinter import filedialog, messagebox
import os
from dotenv import load_dotenv
from tkinter import Toplevel, Label, Entry, Button
import sqlite3
from decimal import Decimal
import shutil
import threading
import smtplib
from email.message import EmailMessage
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.styles import getSampleStyleSheet
from PIL import Image, ImageTk
import pandas as pd
from docxtpl import DocxTemplate
import openpyxl
from docx import Document
from docx2pdf import convert
import re
import win32com.client
import string
from docx.shared import Inches
import unicodedata
import subprocess


"""
load_dotenv()  # încarcă variabilele din .env
API_KEY = os.getenv("API_KEY")
if not API_KEY:
    raise RuntimeError("API KEY lipsă! Verifică fișierul .env")
"""

API_KEY = "p3tze7ux-ft6wflmj-caocdcdc-za3qm51m"  # aici pui cheia api generata altfel ai eroare firma negasita
ADMIN_PASSWORD = "cipri"  # parolă pentru ștergere client din baza de date

DB_NAME = "baza_date.db"


def creeaza_db_si_tabele(db_path):
    """Creează baza de date și tabelele dacă nu există"""
    try:
        conn = sqlite3.connect(db_path)
        cursor = conn.cursor()

        # Tabela clienti
        cursor.execute("""
        CREATE TABLE IF NOT EXISTS tabela_date_clienti (
            Nr_Crt INTEGER PRIMARY KEY AUTOINCREMENT,
            Nume_Firma TEXT,
            Sediu_Social TEXT,
            Cui TEXT UNIQUE,            
            Nr_Telefon TEXT,
            Mail TEXT,
            Reg_Comert TEXT,
            Tva TEXT,
            Administrator TEXT,
            Status_Firma TEXT
        )
        """)

        # Tabela sedii secundare
        cursor.execute("""
        CREATE TABLE IF NOT EXISTS tabela_sedii_secundare (
            Nr_Crt INTEGER PRIMARY KEY AUTOINCREMENT,
            Id_Client INTEGER,
            Punct_Lucru TEXT,
            Status_Punct_Lucru TEXT,
            Model_Amef TEXT,
            Serie_Amef TEXT UNIQUE,
            Nui TEXT UNIQUE,
            Tip_Abonament TEXT,
            Data_Conect_Anaf TEXT,
            Tehnician TEXT,
            Data_Exp_Abon TEXT,
            Val_Ctr REAL,
            Data_Exp_Gprs TEXT,
            Tip_Conect_Anaf TEXT,
            Status_AMEF TEXT,
            Observatii TEXT,
            FOREIGN KEY(Id_Client) REFERENCES tabela_date_clienti(Nr_Crt)
        )
        """)

        cursor.execute("""
        CREATE UNIQUE INDEX IF NOT EXISTS idx_serie_amef
        ON tabela_sedii_secundare(Serie_Amef)
        """)

        cursor.execute("""
        CREATE UNIQUE INDEX IF NOT EXISTS idx_nui
        ON tabela_sedii_secundare(Nui)
        """)

        # Tabela istoric abonamente
        cursor.execute("""
        CREATE TABLE IF NOT EXISTS istoric_abonamente (
            Nr_Crt INTEGER PRIMARY KEY AUTOINCREMENT,
            id_client INTEGER,
            id_sediu INTEGER,
            serie_amef TEXT,
            tip_abonament TEXT,
            data_start TEXT,
            data_expirare TEXT,
            data_prelungire TEXT,
            observatii TEXT
        )
        """)

        conn.commit()
        conn.close()
        print("Baza de date și tabelele au fost create cu succes.")
    except Exception as e:
        print("!!! Eroare la crearea bazei sau a tabelelor:", e)
        messagebox.showerror("Eroare DB", f"Nu am putut crea baza de date: \n{e}")


def conectare_db():
    """Conectează aplicația la baza de date, creează DB dacă nu există"""
    try:
        if getattr(sys, 'frozen', False):
            base_path = os.path.dirname(sys.executable)
        else:
            base_path = os.path.dirname(__file__)

        db_path = os.path.join(base_path, DB_NAME)
        print("Calea bazei de date:", db_path)

        # Dacă baza de date nu există, o creăm
        if not os.path.exists(db_path):
            print("Baza de date nu există, o creez...")
            creeaza_db_si_tabele(db_path)

        # Conectăm la baza existentă
        conn = sqlite3.connect(db_path)
        conn.row_factory = sqlite3.Row
        print("Conectare DB realizată cu succes.")
        return conn

    except Exception as e:
        print("!!! Eroare DB:", e)
        messagebox.showerror("Eroare DB", f"Nu ma pot conecta la baza de date: \n{e}")
        return None


# Exemplu de utilizare
if __name__ == "__main__":
    conn = conectare_db()
    if conn:
        conn.close()


# Functii pentru export/import baza date
def export_db():
    filedialog.asksaveasfilename(
        initialfile="baza_date.db",
        defaultextension=".db"
    )


def import_db():
    filedialog.askopenfilename(filetypes=[("SQLite DB", "*.db")])


"""
Functie pentru a uni 2 baze de date sqlite3 cu verificare si actualizare clienti
si cu creare fisier log cu actualizari
"""

def merge_sqlite_with_file_log(db_source_path, db_target_path, log_file_path="merge_log.txt"):
    # funcție simplă de log în fișier
    def log(msg):
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        with open(log_file_path, "a", encoding="utf-8") as f:
            f.write(f"[{timestamp}] {msg}\n")
        print(msg)  # păstrează și print-ul pe consolă

    def alege_data_mai_buna(data1, data2):
        if not data1:
            return data2
        if not data2:
            return data1

        try:
            if isinstance(data1, str):
                data1 = datetime.fromisoformat(data1)
            if isinstance(data2, str):
                data2 = datetime.fromisoformat(data2)

            return max(data1, data2)
        except:
            return data1

    # conectare baze
    src_conn = sqlite3.connect(db_source_path)
    src_conn.row_factory = sqlite3.Row
    src_cursor = src_conn.cursor()

    tgt_conn = sqlite3.connect(db_target_path)
    tgt_conn.row_factory = sqlite3.Row
    tgt_cursor = tgt_conn.cursor()

    # ================== CREAȚI TABELLELE DACA NU EXISTĂ ==================
    tgt_cursor.execute("""
    CREATE TABLE IF NOT EXISTS tabela_date_clienti (
        Nr_Crt INTEGER PRIMARY KEY AUTOINCREMENT,
        Nume_Firma TEXT,
        Sediu_Social TEXT,
        Cui TEXT UNIQUE,
        Nr_Telefon TEXT,
        Mail TEXT,
        Reg_Comert TEXT,
        Tva TEXT,
        Administrator TEXT,
        Status_Firma TEXT
    )
    """)

    tgt_cursor.execute("""
    CREATE TABLE IF NOT EXISTS tabela_sedii_secundare (
        Nr_Crt INTEGER PRIMARY KEY AUTOINCREMENT,
        Id_Client INTEGER,
        Punct_Lucru TEXT,
        Status_Punct_Lucru TEXT,
        Model_Amef TEXT,
        Serie_Amef TEXT,
        Nui TEXT,
        Tip_Abonament TEXT,
        Data_Conect_Anaf TEXT,
        Tehnician TEXT,
        Data_Exp_Abon TEXT,
        Val_Ctr REAL,
        Data_Exp_Gprs TEXT,
        Tip_Conect_Anaf TEXT,
        Status_AMEF TEXT,
        Observatii TEXT,
        FOREIGN KEY(Id_Client) REFERENCES tabela_date_clienti(Nr_Crt)
    )
    """)

    tgt_cursor.execute("""
    CREATE TABLE IF NOT EXISTS istoric_abonamente (
        Nr_Crt INTEGER PRIMARY KEY AUTOINCREMENT,
        id_client INTEGER,
        id_sediu INTEGER,
        serie_amef TEXT,
        tip_abonament TEXT,
        data_start TEXT,
        data_expirare TEXT,
        data_prelungire TEXT,
        observatii TEXT,
        FOREIGN KEY(id_client) REFERENCES tabela_date_clienti(Nr_Crt),
        FOREIGN KEY(id_sediu) REFERENCES tabela_sedii_secundare(Nr_Crt)
    )
    """)
    tgt_conn.commit()

    log("=== START MERGE BAZE DE DATE ===")

    # ================== MERGE CLIENTI ==================
    src_cursor.execute("SELECT * FROM tabela_date_clienti")
    clienti = src_cursor.fetchall()

    for client in clienti:
        tgt_cursor.execute("SELECT * FROM tabela_date_clienti WHERE Cui=?", (client["Cui"],))
        existing_client = tgt_cursor.fetchone()

        if existing_client:
            updated_fields = []
            updated_values = []
            for field in ["Nume_Firma", "Sediu_Social", "Nr_Telefon", "Mail", "Reg_Comert", "Tva", "Administrator",
                          "Status_Firma"]:
                if client[field] != existing_client[field]:
                    updated_fields.append(f"{field}=?")
                    updated_values.append(client[field])
            if updated_fields:
                updated_values.append(existing_client["Nr_Crt"])
                tgt_cursor.execute(f"UPDATE tabela_date_clienti SET {', '.join(updated_fields)} WHERE Nr_Crt=?",
                                   updated_values)
                log(f"Actualizat client CUI={client['Cui']}: {', '.join(updated_fields)}")
            else:
                log(f"Client CUI={client['Cui']} deja existent, fără modificări")
        else:
            fields = ["Nume_Firma", "Sediu_Social", "Cui", "Nr_Telefon", "Mail", "Reg_Comert", "Tva", "Administrator",
                      "Status_Firma"]
            placeholders = ",".join("?" * len(fields))
            values = [client[f] for f in fields]
            tgt_cursor.execute(f"INSERT INTO tabela_date_clienti ({','.join(fields)}) VALUES ({placeholders})", values)
            log(f"Adăugat client nou: CUI={client['Cui']}")

    tgt_conn.commit()

    # ================== MERGE SEDII SECUNDARE ==================
    src_cursor.execute("SELECT * FROM tabela_sedii_secundare")
    sedii = src_cursor.fetchall()

    for sediu in sedii:
        src_cursor.execute("SELECT Cui FROM tabela_date_clienti WHERE Nr_Crt=?", (sediu["Id_Client"],))
        cui = src_cursor.fetchone()["Cui"]

        tgt_cursor.execute("SELECT Nr_Crt FROM tabela_date_clienti WHERE Cui=?", (cui,))
        target_client = tgt_cursor.fetchone()
        if not target_client:
            log(f"Client CUI={cui} nu există în baza țintă, se sare sediul Serie_Amef={sediu['Serie_Amef']}")
            continue
        id_client_target = target_client["Nr_Crt"]

        tgt_cursor.execute("""
            SELECT * FROM tabela_sedii_secundare
            WHERE Id_Client=? AND Serie_Amef=? AND Nui=?
        """, (id_client_target, sediu["Serie_Amef"], sediu["Nui"]))
        existing_sediu = tgt_cursor.fetchone()

        if existing_sediu:
            updated_fields = []
            updated_values = []
            for field in ["Punct_Lucru", "Status_Punct_Lucru", "Model_Amef", "Tip_Abonament", "Data_Conect_Anaf", "Tehnician",
                          "Data_Exp_Abon", "Val_Ctr", "Data_Exp_Gprs", "Tip_Conect_Anaf", "Status_AMEF", "Observatii"]:
                val = sediu[field]
                existing_val = existing_sediu[field]

                if isinstance(val, Decimal):
                    val = float(val)
                # Sectiune care protejeaza baza de date la suprascriere import date vechi peste cele mai noi
                if field in ["Data_Exp_Abon", "Data_Exp_Gprs"]:
                    val_final = alege_data_mai_buna(existing_val, val)

                    if val_final != existing_val:
                        updated_fields.append(f"{field}=?")
                        updated_values.append(val_final)
                        log(f"{field} protejat: {existing_val} → {val_final}")

                else:
                    if val != existing_sediu[field]:
                        updated_fields.append(f"{field}=?")
                        updated_values.append(val)

            if updated_fields:
                updated_values.append(existing_sediu["Nr_Crt"])
                tgt_cursor.execute(f"UPDATE tabela_sedii_secundare SET {', '.join(updated_fields)} WHERE Nr_Crt=?",
                                   updated_values)
                log(f"Actualizat sediu Serie_Amef={sediu['Serie_Amef']} client CUI={cui}: {', '.join(updated_fields)}")
            else:
                log(f"Sediu Serie_Amef={sediu['Serie_Amef']} client CUI={cui} deja existent, fără modificări")
        else:
            fields = ["Id_Client", "Punct_Lucru", "Status_Punct_Lucru", "Model_Amef", "Serie_Amef", "Nui", "Tip_Abonament",
                      "Data_Conect_Anaf", "Tehnician", "Data_Exp_Abon", "Val_Ctr", "Data_Exp_Gprs", "Tip_Conect_Anaf", "Status_AMEF", "Observatii"]
            values = [id_client_target] + [float(sediu[f]) if isinstance(sediu[f], Decimal) else sediu[f] for f in
                                           fields[1:]]
            placeholders = ",".join("?" * len(fields))
            tgt_cursor.execute(f"INSERT INTO tabela_sedii_secundare ({','.join(fields)}) VALUES ({placeholders})",
                               values)
            log(f"Adăugat sediu nou Serie_Amef={sediu['Serie_Amef']} client CUI={cui}")

    tgt_conn.commit()

    # ================== MERGE ISTORIC_ABONAMENTE ==================
    src_cursor.execute("SELECT * FROM istoric_abonamente")
    istoric = src_cursor.fetchall()

    for entry in istoric:
        src_cursor.execute("SELECT Cui FROM tabela_date_clienti WHERE Nr_Crt=?", (entry["id_client"],))
        cui = src_cursor.fetchone()["Cui"]

        tgt_cursor.execute("SELECT Nr_Crt FROM tabela_date_clienti WHERE Cui=?", (cui,))
        target_client = tgt_cursor.fetchone()
        if not target_client:
            log(f"Client CUI={cui} nu există în baza țintă, se sare istoricul Serie_Amef={entry['serie_amef']}")
            continue
        id_client_target = target_client["Nr_Crt"]

        tgt_cursor.execute("""
            SELECT Nr_Crt FROM tabela_sedii_secundare 
            WHERE Id_Client=? AND Serie_Amef=?
        """, (id_client_target, entry["serie_amef"]))

        target_sediu = tgt_cursor.fetchone()
        id_sediu_target = target_sediu["Nr_Crt"] if target_sediu else None

        if not id_sediu_target:
            log(f"Sediu lipsă în target pentru Serie_Amef={entry['serie_amef']}")
            continue

        tgt_cursor.execute("""
            SELECT * FROM istoric_abonamente 
            WHERE id_sediu=? AND tip_abonament=? AND data_start=?
        """, (id_sediu_target, entry["tip_abonament"], entry["data_start"]))
        existing_entry = tgt_cursor.fetchone()

        if existing_entry:
            updated_fields = []
            updated_values = []
            for field in ["data_expirare", "observatii"]:
                if entry[field] != existing_entry[field]:
                    updated_fields.append(f"{field}=?")
                    updated_values.append(entry[field])
            if updated_fields:
                updated_values.append(existing_entry["Nr_Crt"])
                tgt_cursor.execute(f"UPDATE istoric_abonamente SET {', '.join(updated_fields)} WHERE Nr_Crt=?",
                                   updated_values)
                log(f"Actualizat istoric abonament Serie_Amef={entry['serie_amef']} client CUI={cui}: {', '.join(updated_fields)}")
            else:
                log(f"Istoric abonament Serie_Amef={entry['serie_amef']} client CUI={cui} deja existent, fără modificări")
        else:
            fields = ["id_client", "id_sediu", "serie_amef", "tip_abonament", "data_start", "data_expirare",
                      "observatii"]
            values = [id_client_target, id_sediu_target] + [entry[f] for f in fields[2:]]
            placeholders = ",".join("?" * len(fields))
            tgt_cursor.execute(f"INSERT INTO istoric_abonamente ({','.join(fields)}) VALUES ({placeholders})", values)
            log(f"Adăugat istoric abonament Serie_Amef={entry['serie_amef']} client CUI={cui}")

    tgt_conn.commit()
    src_conn.close()
    tgt_conn.close()
    log("=== MERGE COMPLET FINALIZAT ===")

    def buton_update_baza(db_source_path, db_target_path):
        # cerere parola
        parola = simpledialog.askstring("Parola Admin", "Introduceți parola pentru update:", show="*")
        if parola != ADMIN_PASSWORD:
            messagebox.showerror("Eroare", "Parola incorectă!")
            return

        # confirmare
        if not messagebox.askyesno("Confirmare", "Sigur doriți să faceți update și merge la baza de date?"):
            return

        try:
            merge_sqlite_with_file_log(db_source_path, db_target_path)
            messagebox.showinfo("Succes", f"Update și merge complet realizat!\nVezi detalii în 'merge_log.txt'")
        except Exception as e:
            messagebox.showerror("Eroare", f"A apărut o eroare la merge: {e}")


# Final functie merge baze date

def update_baza_protejat():
    parola = simpledialog.askstring("Parola Admin", "Introduceți parola pentru update:", show="*")
    if parola != ADMIN_PASSWORD:
        messagebox.showerror("Eroare", "Parola incorectă!")
        return

    # Alegem baza de date sursa
    source = filedialog.askopenfilename(
        title="Selecteaza baza sursa (din care importi datele)",
        filetypes=[("SQLite DB", "*.db")]
    )
    if not source:
        return

    # Alegem baza de date tinta
    target = filedialog.askopenfilename(
        title="Selecteaza baza tinta (baza date master)",
        filetypes=[("SQLite DB", "*.db")]
    )
    if not source:
        return

    if source == target:
        messagebox.showerror("Eroare", "Nu poti selecta aceeasi baza")
        return

    # === BACKUP ===
    try:
        backup_file = backup_database(target)
    except Exception as e:
        messagebox.showerror("Eroare Backup", str(e))
        return

    # === Fereastră progres ===
    progress_win = Toplevel()
    progress_win.title("Merge în curs...")
    progress = ttk.Progressbar(progress_win, mode="indeterminate", length=400)
    progress.pack(padx=20, pady=20)
    progress["value"] = 0

    def run_merge():
        try:
            merge_sqlite_with_file_log(source, target)

            def finish_merge():
                progress.stop()
                progress_win.destroy()
                messagebox.showinfo("Succes", f"Merge finalizat!\nBackup salvat:\n{backup_file}")

                # ---------------- Întrebare trimitere email ----------------
                if messagebox.askyesno("Email", "Doriți să trimiteți backup-ul pe email?"):
                    destinatar = simpledialog.askstring("Email destinatar", "Introduceți adresa de email:")
                    if destinatar:
                        trimite_backup_email(backup_file, destinatar)
                        messagebox.showinfo("Email trimis", f"Backup-ul a fost trimis către {destinatar}")

            progress_win.after(0, finish_merge)

        except Exception as e:
            progress.stop()
            progress_win.destroy()
            messagebox.showerror("Eroare", str(e))

    threading.Thread(target=run_merge, daemon=True).start()


# Functie pentru backup baza date inainte de merge
def backup_database(db_path):
    if not os.path.exists(db_path):
        raise FileNotFoundError("Baza de date nu exista !")

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_name = f"{os.path.splitext(db_path)[0]}_backup_{timestamp}.db"
    shutil.copy2(db_path, backup_name)
    return backup_name
# Final functie backup baz adate


# Functie trimitere backup baza date pe mail
def trimite_backup_email(backup_path, destinatar):
    msg = EmailMessage()
    msg["Subject"] = "Backup baza de date clienti SQLite3"
    msg["From"] = "instalari.secretdata@gmail.com"
    msg["To"] = destinatar
    msg.set_content("Atasat gasiti backup-ul bazei de date.")

    with open(backup_path, "rb") as f:
        file_data = f.read()
        file_name = os.path.basename(backup_path)
    msg.add_attachment(file_data, maintype="application", subtype="octet-stream", filename=file_name)

    with smtplib.SMTP_SSL("smtp.gmail.com", 465) as smtp:
        smtp.login("instalari.secretdata@gmail.com", "ipns zunb qyxe bqbl")
        smtp.send_message(msg)


# Functie cautare cu api anaf
# -------------------------
# Normalizare nume firma
# -------------------------


def normalize_nume_firma(nume):
    if not nume:
        return ""

    # Eliminare diacritice
    traducere = str.maketrans({
        "ă": "a",
        "â": "a",
        "î": "i",
        "ș": "s",
        "ş": "s",
        "ț": "t",
        "ţ": "t",
        "Ă": "A",
        "Â": "A",
        "Î": "I",
        "Ș": "S",
        "Ş": "S",
        "Ț": "T",
        "Ţ": "T"
    })

    nume = nume.translate(traducere)
    nume = nume.strip().lower()

    replacements = {
        "persoana fizica autorizata": "PFA",
        "intreprindere individuala": "II",
        "intreprindere familiala": "IF",
        " srl": " SRL",
        " sa": " SA",
        " s.r.l.": " SRL",
        "sc ": "SC "
    }

    for k, v in replacements.items():
        nume = nume.replace(k, v)

    nume = nume.strip()

    if (" SRL" in nume or " SA" in nume) and not nume.startswith("SC "):
        nume = "SC " + nume

    return nume.upper()


# -------------------------
# Parsare adresa (din string ANAF)
# -------------------------
def parse_adresa_anaf(adresa):
    parts = adresa.split(",")

    judet = ""
    localitate = ""
    strada = "Principala"
    numar = "-"
    bloc = ""
    scara = ""
    apartament = ""

    for p in parts:
        p = p.strip().lower()

        if "jud." in p:
            judet = p.replace("jud.", "").strip()

        elif "mun." in p or "oras" in p or "loc." in p:
            localitate = (
                p.replace("mun.", "")
                 .replace("oras", "")
                 .replace("loc.", "")
                 .strip()
            )

        elif "sat" in p:
            localitate = p.replace("sat", "").strip()

        elif "str." in p:
            strada = p.replace("str.", "").strip()

        elif "nr." in p:
            numar = p.replace("nr.", "").strip().upper()

        elif "bl." in p:
            bloc = p.replace("bl.", "").strip().upper()

        elif "sc." in p:
            scara = p.replace("sc.", "").strip().upper()

        elif "ap." in p:
            apartament = p.replace("ap.", "").strip().upper()

    return {
        "judet": judet.title(),
        "localitate": localitate.title(),
        "strada": strada.title(),
        "numar": numar,
        "bloc": bloc,
        "scara": scara,
        "apartament": apartament
    }


# =========================
# FUNCTIA PRINCIPALA (ANAF)
# =========================
def cauta_firma_firmeapi(cui):
    cui = cui.strip().replace("RO", "").replace("ro", "")

    url = "https://webservicesp.anaf.ro/api/PlatitorTvaRest/v9/tva"

    payload = [{
        "cui": int(cui),
        "data": datetime.now().strftime("%Y-%m-%d")
    }]

    try:
        r = requests.post(url, json=payload, timeout=5)
        r.raise_for_status()
    except requests.RequestException:
        return None

    data = r.json()

    # -------------------------
    # CUI invalid
    # -------------------------
    if data.get("notFound"):
        return {
            "cui": "CUI INVALID",
            "nume": "",
            "adresa": "",
            "reg_comert": "",
            "tva": "NU"
        }

    found = data.get("found", [])
    if not found:
        return None

    firma = found[0]

    date_generale = firma.get("date_generale", {})
    tva_info = firma.get("inregistrare_scop_Tva", {})

    # TVA
    este_tva = tva_info.get("scpTVA", False)
    tva_text = "DA" if este_tva else "NU"

    # DATE FIRMA
    cui_return = str(date_generale.get("cui", cui))
    if este_tva:
        cui_return = f"RO{cui_return}"

    nume = normalize_nume_firma(date_generale.get("denumire", ""))
    adresa = date_generale.get("adresa", "")
    reg_com = date_generale.get("nrRegCom", "")

    # PARSARE ADRESA
    adresa_parsata = parse_adresa_anaf(adresa)

    # OUTPUT FINAL

    return {
        "cui": cui_return,
        "nume": nume,
        "adresa": adresa,
        "reg_comert": reg_com,
        "tva": tva_text,

        "judet": adresa_parsata["judet"],
        "localitate": adresa_parsata["localitate"],
        "strada": adresa_parsata["strada"],
        "numar": adresa_parsata["numar"],
        "bloc": adresa_parsata["bloc"],
        "scara": adresa_parsata["scara"],
        "apartament": adresa_parsata["apartament"],
    }
# final functie cautare cu api anaf



"""
decomentam functia pentru utilizarea cu firmeapi.ro
# =========================
# Functie de cautare firma in api dupa cod fiscal client
# =========================
def cauta_firma_firmeapi(cui):
    cui = cui.strip().replace("RO", "").replace("ro", "")
    url = f"https://www.firmeapi.ro/api/v1/firma/{cui}"
    headers = {"X-API-KEY": API_KEY, "Accept": "application/json"}
    try:
        r = requests.get(url, headers=headers, timeout=10)
        r.raise_for_status()
    except requests.RequestException:
        return None
    firma = r.json().get("data")
    if not firma:
        return None
    adresa_completa = ""
    sediu = firma.get("adresa_sediu_social", {})
    if isinstance(sediu, dict):
        strada = sediu.get("strada", "")
        numar = sediu.get("numar", "")
        localitate = ""
        judet = ""
        if isinstance(sediu.get("localitate"), dict):
            local = sediu["localitate"]
            localitate = local.get("nume", "")
            judet = local.get("judet", {}).get("nume", "")
        elif isinstance(sediu.get("localitate"), str):
            localitate = sediu.get("localitate")
        adresa_completa = f"{strada} {numar}, {localitate}, {judet}".strip(" ,")
    return {
        "cui": firma.get("cui", cui),
        "nume": firma.get("denumire", ""),
        "adresa": adresa_completa,
        "reg_comert": firma.get("nr_reg_com", "")
    }
"""



"""
 Zona de FUNCȚII CRUD
 functie de conectare la baza de date sqlite date_clienti care are 3 tabele:
 tabela_date_clienti, tabela_sedii_secundare si tabela istoric_abonamente

"""


# Functia asta nu mai e utilizata in cod
def incarca_dropdown_puncte():
    for i in tree.get_children():
        tree.delete(i)
    conn = conectare_db()
    cursor = conn.cursor()
    cursor.execute("""SELECT d.Nr_Crt, d.Nume_Firma, d.Cui, d.Sediu_Social,
                      s.Punct_Lucru, s.Model_Amef, s.Serie_Amef
                      FROM tabela_date_clienti d
                      LEFT JOIN tabela_sedii_secundare s ON d.Nr_Crt = s.Id_Client""")
    for row in cursor.fetchall():
        tree.insert("", "end", values=row)
    conn.close()


# Functie pentru golirea tuturor campurilor din interfata
def resetare_toate_campurile():
    for dic in (entries_client, entries_sediu):
        for widget in dic.values():
            # Entry
            try:
                widget.delete(0, tk.END)
            except:
                pass
            # Combobox readonly
            try:
                if isinstance(widget, ttk.Combobox):
                    state_orig = widget.cget("state")
                    widget.config(state="normal")
                    widget.set("")
                    widget.config(state=state_orig)
            except:
                pass
            # DateEntry
            try:
                widget.set_date("")
            except:
                pass

        # Variabile live pentru NUI și Serie AMEF
    entry_nui_var.set("")
    entry_serie_amef_var.set("")

    entry_cui.focus()


# functie de resetare a campului de cautare client
def resetare_camp_cautare():
    search_entry.delete(0, tk.END)
    for item in tree.get_children():
        tree.delete(item)


# Funtie pentru modificarea datelor introduse gresit
"""Populează câmpurile cu datele clientului după CUI și punct de lucru pentru editare
Nu o mai folosesc in cod pentru ca si salvare client face acelasi lucru,
o pastrez, am dezactivat butonul atasat functiei
"""


# def modifica_date_client():
#     cui = entry_cui.get().strip()
#     serie_amef = entry_serie_amef.get().strip()
#
#     if not cui:
#         messagebox.showwarning("Eroare", "Introduceți CUI-ul clientului")
#         return
#
#     conn = conectare_db()
#     cursor = conn.cursor()
#
#     cursor.execute("""
#             SELECT d.Nume_Firma, d.Cui, d.Reg_Comert, d.Tva, d.Sediu_Social,
#                    s.Punct_Lucru, s.Model_Amef, s.Serie_Amef, s.Nui,
#                    s.Tehnician, s.Data_Conect_Anaf, s.Data_Exp_Abon, s.Val_Ctr, s.Data_Exp_Gprs
#             FROM tabela_date_clienti d
#             LEFT JOIN tabela_sedii_secundare s ON d.Nr_Crt = s.Id_Client
#             WHERE d.Cui=? AND s.Serie_Amef=?
#         """, (cui, serie_amef))
#     result = cursor.fetchone()
#     conn.close()
#
#     if not result:
#         messagebox.showinfo("Info", "Nu s-a găsit clientul sau punctul de lucru")
#         return
#
#     # Populare câmpuri
#     mapping = {
#         "Nume firmă": result["Nume_Firma"],
#         "CUI Client": result["Cui"],
#         "Nr. Registrul Comertului": result["Reg_Comert"],
#         "Plătitor TVA": result["Tva"],
#         "Adresă sediu": result["Sediu_Social"],
#         "Punct de lucru": result["Punct_Lucru"],
#         "Model Amef": result["Model_Amef"],
#         "Serie Amef": result["Serie_Amef"],
#         "Nui Amef": result["Nui"],
#         "Tehnician Service": result["Tehnician"],
#         "Data conectare Anaf": result["Data_Conect_Anaf"],
#         "Data expirare abonament": result["Data_Exp_Abon"],
#         "Valoare contract - RON": result["Val_Ctr"],
#         "Data expirare gprs": result["Data_Exp_Gprs"]
#     }
#
#     for label, value in mapping.items():
#         entries[label].delete(0, tk.END)
#         entries[label].insert(0, value)


def cauta_firma():
    cui = entry_cui.get().strip()
    if not cui:
        messagebox.showwarning("Eroare", "Introduceți un CUI")
        return
    info = cauta_firma_firmeapi(cui)
    if not info:
        messagebox.showinfo("Info", "Firma nu a fost găsită")
        return

    # -------------------------
    # TVA + culoare
    # -------------------------
    if info["tva"] == "DA":
        label_tva.config(text="TVA: DA", fg="green")
        culoare_tva = "DA"
    else:
        label_tva.config(text="TVA: NU", fg="red")
        culoare_tva = "NU"



    # Logica adaugare RO sau fara Ro in functie de TVA
    cui_final = str(info["cui"])
    if info["tva"] == "DA":
        if not cui_final.upper().startswith("RO"):
            cui_final = "RO" + cui_final
    else:
        cui_final = cui_final.replace("RO", "").replace("ro", "")
    entry_cui.delete(0, tk.END)
    entry_cui.insert(0, cui_final)


    entry_nume.delete(0, tk.END)
    entry_nume.insert(0, info["nume"])
    entry_adresa.delete(0, tk.END)
    entry_adresa.insert(0, info["adresa"])
    entry_reg_comert.delete(0, tk.END)
    entry_reg_comert.insert(0, info["reg_comert"])
    entry_tva.delete(0, tk.END)
    entry_tva.insert(0, info["tva"])


"""
# Functie pentru introducerea clientilor in baza de date dar se si poate modifica datele clientului, 
La salvarea data expirare service si data expirare gprs , cu aceasta functie nu se salveaza in istoric abonamente 
decat cu butonul de prelungire abonament
"""

# functie dde verificare unicitate serie amef si nui amef
def verifica_serie_nui_unice(serie, nui, nr_crt_sediu=None):
    conn = conectare_db()
    cursor = conn.cursor()

    cursor.execute("""
        SELECT s.Nr_Crt, d.Nume_Firma, s.Serie_Amef, s.Nui
        FROM tabela_sedii_secundare s
        JOIN tabela_date_clienti d
        ON d.Nr_Crt = s.Id_Client
        WHERE (s.Serie_Amef=? OR s.Nui=?)
    """, (serie, nui))
    rows = cursor.fetchall()
    conn.close()

    for r in rows:
        # permite editarea la acelasi client
        if nr_crt_sediu and r["Nr_Crt"] == nr_crt_sediu:
            continue
        return  r
    return None

# Validare NUI: doar cifre, max 10, paste blocat dacă depășește
def validare_nui(text_nou):
    if text_nou == "":
        return True
    if not text_nou.isdigit():
        return False
    if len(text_nou) > 10:
        return False
    return True

# Validare Serie AMEF: litere + cifre, uppercase automat, fără spații și caractere speciale, max 20 caractere
def validare_serie_amef(text_nou):
    if text_nou == "":
        return True

    # elimină spații și caractere speciale
    valid_chars = string.ascii_letters + string.digits
    filtrat = "".join(c for c in text_nou if c in valid_chars).upper()

    # dacă lungimea depășește 20, taie
    if len(filtrat) > 20:
        filtrat = filtrat[:20]

    # actualizează Entry cu text curățat
    entry_serie_amef_var.set(filtrat)

    return True


def salveaza_client():
    data = {
        "cui": entry_cui.get().strip(),
        "nume": entry_nume.get().strip(),
        "adresa": entry_adresa.get().strip(),
        "reg_comert": entry_reg_comert.get().strip(),
        "tva": entry_tva.get().strip(),
        "administrator": entry_administrator.get().strip(),
        "status_firma": entry_status_firma.get().strip(),
        "telefon": entry_telefon.get().strip(),
        "mail": entry_mail.get().strip(),
        "punct_lucru": entry_punct_lucru.get().strip(),
        "status_punct_lucru": entry_status_punct_lucru.get().strip(),
        "model_amef": entry_model_amef.get().strip(),
        "serie_amef": entry_serie_amef.get().strip(),
        "nui": entry_nui.get().strip(),
        "tip_abonament": entry_tip_abonament.get().strip(),
        "data_conect": entry_conectare_anaf.get().strip(),
        "tehnician": entry_tehnician.get().strip(),
        "data_exp": entry_data_exp.get().strip(),
        "val_ctr": entry_val_ctr.get().strip(),
        "data_exp_gprs": entry_data_exp_gprs.get().strip(),
        "tip_conect_anaf": entry_tip_conect_anaf.get().strip(),
        "status_amef": entry_status_amef.get().strip(),
        "observatii": entry_observatii.get().strip()
    }

    if not data["cui"] or not data["nume"]:
        messagebox.showwarning("Eroare", "CUI și Nume Firmă sunt obligatorii!")
        return

    if not data["serie_amef"]:
        messagebox.showwarning("Eroare", "Seria AMEF este obligatorie!")
        return

    if data["nui"]:
        # aici verificam daca NUI este din 10 cifre
        if not data["nui"].isdigit() or len(data["nui"]) !=10:
            messagebox.showerror(
                "Eroare NUI",
                "NUI trebuie sa contina exact 10 cifre"
            )
            return

    conn = conectare_db()
    if not conn:
        return
    cursor = conn.cursor()

    # verific client existent
    try:
        cursor.execute("SELECT Nr_Crt FROM tabela_date_clienti WHERE Cui=?",
                       (data["cui"],))
        result = cursor.fetchone()

        if result:
            id_client = result["Nr_Crt"]
            cursor.execute("""
                UPDATE tabela_date_clienti
                SET Nume_Firma=?, Sediu_Social=?, Tva=?, Administrator=?,
                    Status_Firma=?, Nr_Telefon=?, Mail=?
                WHERE Nr_Crt=?
            """, (
                data["nume"], data["adresa"], data["tva"], data["administrator"],
                data["status_firma"], data["telefon"], data["mail"], id_client
            ))
            messagebox.showinfo("Info", f"Client existent. Datele au fost actualizate (Nr_Crt={id_client})")
        else:
            cursor.execute("""
                INSERT INTO tabela_date_clienti
                (Nume_Firma, Sediu_Social, Cui, Nr_Telefon, Mail, Reg_Comert, Tva, Administrator, Status_Firma)
                VALUES (?,?,?,?,?,?,?,?,?)
            """, (
                data["nume"], data["adresa"], data["cui"], data["telefon"], data["mail"],
                data["reg_comert"], data["tva"], data["administrator"], data["status_firma"]
            ))
            id_client = cursor.lastrowid
            messagebox.showinfo("Succes", f"Client nou adăugat (Nr_Crt={id_client})")

        # punct de lucru
        cursor.execute("SELECT Nr_Crt FROM tabela_sedii_secundare WHERE Id_Client=? AND Serie_Amef=?",
                       (id_client, data["serie_amef"]))
        row_sediu = cursor.fetchone()
        nr_crt_sediu = None
        if row_sediu:
            nr_crt_sediu = row_sediu["Nr_Crt"]

        # ==============================
        # verificare duplicat serie / nui
        # ==============================

        rez = verifica_serie_nui_unice(
            data["serie_amef"],
            data["nui"],
            nr_crt_sediu
        )

        if rez:
            messagebox.showerror(
                "Duplicat",
                f"Seria sau NUI există deja!\n\n"
                f"Firma: {rez['Nume_Firma']}\n"
                f"Serie: {rez['Serie_Amef']}\n"
                f"NUI: {rez['Nui']}"
            )

            conn.close()
            return

        # ==============================
        # UPDATE sau INSERT sediu
        # ==============================

        if row_sediu:

            cursor.execute("""
                UPDATE tabela_sedii_secundare
                SET Punct_Lucru=?, Status_Punct_Lucru=?,  Model_Amef=?, Nui=?, Tip_Abonament=?,
                    Data_Conect_Anaf=?, Tehnician=?, Data_Exp_Abon=?,
                    Val_Ctr=?, Data_Exp_Gprs=?, Tip_Conect_Anaf=?, Status_AMEF=?, Observatii=?
                WHERE Nr_Crt=?
            """, (
                data["punct_lucru"], data["status_punct_lucru"], data["model_amef"], data["nui"], data["tip_abonament"],
                data["data_conect"], data["tehnician"], data["data_exp"],
                data["val_ctr"], data["data_exp_gprs"], data["tip_conect_anaf"], data["status_amef"], data["observatii"], nr_crt_sediu
            ))

            messagebox.showinfo(
                "Succes",
                f"Punct de lucru {data['serie_amef']} actualizat"
            )
        else:
            cursor.execute("""
                INSERT INTO tabela_sedii_secundare
                (Id_Client, Punct_Lucru, Status_Punct_Lucru, Model_Amef, Serie_Amef, Nui,
                 Tip_Abonament, Data_Conect_Anaf, Tehnician, Data_Exp_Abon, Val_Ctr, Data_Exp_Gprs, Tip_Conect_Anaf,Status_AMEF, Observatii)
                VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
            """, (
                id_client, data["punct_lucru"], data["status_punct_lucru"], data["model_amef"], data["serie_amef"], data["nui"],
                data["tip_abonament"], data["data_conect"], data["tehnician"],
                data["data_exp"], data["val_ctr"], data["data_exp_gprs"], data["tip_conect_anaf"], data["status_amef"], data["observatii"]
            ))
            messagebox.showinfo("Succes", f"Punct de lucru {data['serie_amef']} adăugat")

        conn.commit()
    except Exception as e:
        conn.rollback()
        messagebox.showerror("Eroare DB", str(e))
    finally:
        conn.close()


"""
Functie penntru calcularea automata a contractului in functie de situatia clientului
platitor/neplatitor tva sau deplasare/anual
"""


def calculeaza_valoare_contract(tip_abonament, platitor_tva):
    TVA = 0.21  # Aici modifici cand se schimba tva-ul
    tip_abonament = tip_abonament.strip().upper()
    platitor_tva = platitor_tva.strip().upper()
    valori_baza = {
        "DEPLASARE-INTERN": 120,
        "DEPLASARE-EXTERN": 135,
        "ANUAL": 300
    }
    if tip_abonament not in valori_baza:
        return ""
    valoare = valori_baza[tip_abonament]
    if platitor_tva == "DA":
        valoare *= (1 + TVA)
    return f"{valoare:.2f}"


# Functie pentru actualizare automata a campului UI
def actualizeaza_valoare_contract(event=None):
    tip = entry_tip_abonament.get()
    tva = entry_tva.get()
    valoare = calculeaza_valoare_contract(tip, tva)
    entry_val_ctr.delete(0, tk.END)
    entry_val_ctr.insert(0, valoare)


"""
Functie pentru a sterge un client din baza de date 
ATENTIE: La stergerea unui client se va sterge toate punctele de lucru si casele de marcat ale clientului
"""


def sterge_client():
    parola = simpledialog.askstring("Parola Admin", "Introduceți parola pentru ștergere:", show="*")
    if parola != ADMIN_PASSWORD:
        messagebox.showerror("Eroare", "Parola incorectă!")
        return
    cui = entry_cui.get().strip()
    if not cui:
        messagebox.showwarning("Eroare", "Introduceți CUI-ul clientului")
        return
    if not messagebox.askyesno("Confirmare", f"Sigur doriți să ștergeți clientul {cui} și toate punctele sale?"):
        return

    conn = conectare_db()
    if not conn:
        return
    cursor = conn.cursor()

    try:
        cursor.execute("SELECT Nr_Crt FROM tabela_date_clienti WHERE Cui=?", (cui,))
        result = cursor.fetchone()
        if not result:
            messagebox.showinfo("Info", "Clientul nu există")
            conn.close()
            return
        id_client = result["Nr_Crt"]
        cursor.execute("DELETE FROM tabela_sedii_secundare WHERE Id_Client=?", (id_client,))
        cursor.execute("DELETE FROM tabela_date_clienti WHERE Nr_Crt=?", (id_client,))
        conn.commit()
        conn.close()
        messagebox.showinfo("Succes", f"Client {cui} și punctele sale au fost șterse")
        resetare_toate_campurile()
        incarca_dropdown_puncte()
    except Exception as e:
        conn.rollback()
        messagebox.showerror("Eroare DB", str(e))

    finally:
        conn.close()


"""
Functie pentru a sterge doar punctul de lucru al clientului, daca se inchide punctul de lucru
Nu se sterge si clientul din baza de date
"""


def sterge_punct():
    parola = simpledialog.askstring("Parola Admin", "Introduceți parola pentru ștergere:", show="*")
    if parola != ADMIN_PASSWORD:
        messagebox.showerror("Eroare", "Parola incorectă!")
        return
    serie_amef = entry_serie_amef.get().strip()
    if not serie_amef:
        messagebox.showwarning("Eroare", "Introduceți seria AMEF a punctului")
        return
    if not messagebox.askyesno("Confirmare", f"Sigur doriți să ștergeți punctul {serie_amef}?"):
        return
    conn = conectare_db()
    if not conn:
        return
    cursor = conn.cursor()
    try:
        cursor.execute("DELETE FROM tabela_sedii_secundare WHERE Serie_Amef=?", (serie_amef,))
        conn.commit()
        conn.close()
        messagebox.showinfo("Succes", f"Punct {serie_amef} a fost șters")
        resetare_toate_campurile()
        incarca_dropdown_puncte()
    except Exception as e:
        conn.rollback()
        messagebox.showerror("Eroare DB", str(e))

    finally:
        conn.close()


"""
Zona de cautare a unui client in baza de date
Functie pentru cautare client in baza de date, dupa cui, nume, serie casa sau nui
"""


def cauta_in_treeview():
    query = search_entry.get().strip().lower()

    # Curățare Treeview
    for item in tree.get_children():
        tree.delete(item)
    found = False  # Flag pentru a vedea daca am gasit la cautare ceva

    conn = conectare_db()
    if not conn:
        return
    cursor = conn.cursor()
    try:
        cursor.execute("""
            SELECT 
                d.Nr_Crt,
                d.Nume_Firma,
                d.Cui,
                d.Sediu_Social,
                d.Nr_Telefon,
                d.Mail,
                d.Reg_Comert,
                d.Tva,
                d.Administrator,
                d.Status_Firma,
                s.Punct_Lucru,
                s.Status_Punct_Lucru,
                s.Model_Amef,
                s.Serie_Amef,
                s.Nui,
                s.Tehnician,
                s.Data_Conect_Anaf,
                s.Data_Exp_Abon,
                s.Val_Ctr,
                s.Tip_Abonament,
                s.Data_Exp_Gprs,
                s.Tip_Conect_Anaf,
                s.Status_AMEF,
                s.Observatii

            FROM tabela_date_clienti d
            LEFT JOIN tabela_sedii_secundare s ON d.Nr_Crt = s.Id_Client
        """)

        rows = cursor.fetchall()
    except Exception as e:
        messagebox.showerror("Eroare DB", str(e))
        return
    finally:
        conn.close()

    for row in rows:
        # protecție la None
        nume = (row["Nume_Firma"] or "").lower()
        cui = (row["Cui"] or "").lower()
        serie = (row["Serie_Amef"] or "").lower()
        nui = (row["Nui"] or "").lower()

        if (
                query in str(row["Nume_Firma"]).lower()
                or query in str(row["Cui"]).lower()
                or query in str(row["Serie_Amef"]).lower()
                or query in str(row["Nui"]).lower()
        ):

            # Tag special pentru status firma non-activ
            status_firma = (row["Status_Firma"] or "").strip().upper()
            status_punct = (row["Status_Punct_Lucru"] or "").strip().upper()

            firma_inactiva = status_firma in [
                "INCHIS",
                "SUSPENDAT",
                "INACTIV-RENUNTAT"
            ]

            punct_inactiv = status_punct in [
                "INCHIS",
                "SUSPENDAT",
                "INACTIV-RENUNTAT"
            ]

            # daca firma este inchisa SAU punctul este inchis
            if firma_inactiva or punct_inactiv:
                tag_final = "status_inactiv"  # Culoarea in tree a firmei inactiva
            else:
                tag_amef = calculeaza_tag_abonament(row["Data_Exp_Abon"])
                tag_gprs = calculeaza_tag_abonament_gprs(row["Data_Exp_Gprs"])
                tag_final = combina_taguri(tag_amef, tag_gprs)
            # inserare rând în Treeview cu tagurile corecte
            tree.insert("", "end", values=(
                row["Nr_Crt"],
                row["Nume_Firma"],
                row["Cui"],
                row["Sediu_Social"],
                row["Nr_Telefon"],
                row["Mail"],
                row["Reg_Comert"],
                row["Tva"],
                row["Administrator"],
                row["Status_Firma"],
                row["Punct_Lucru"],
                row["Status_Punct_Lucru"],
                row["Model_Amef"],
                row["Serie_Amef"],
                row["Nui"],
                row["Tehnician"],
                row["Data_Conect_Anaf"],
                row["Data_Exp_Abon"],
                row["Val_Ctr"],
                row["Tip_Abonament"],
                row["Data_Exp_Gprs"],
                row["Tip_Conect_Anaf"],
                row["Status_AMEF"],
                row["Observatii"]

            ),
                        tags=(tag_final,)
                        )
            print("STATUS DIN DB =", repr(row["Status_Firma"]))
            found = True  # Am gasit la cautare ceva
    # Daca nu am gasit nimic la cautare afiseaza nu am gait nimic in baza de date
    if not found:
        messagebox.showinfo("Rezultat cautare", f"Nici o inregistrare nu a fost gasita pentru: {query}")


"""
Functie in care combinam cele 2 taguri de amef si gprs
pentru colorarea coloanelor din tree cautare 
culoare rosie daca oricare din abonamente este expirat 
culoare galbena daca oricare din abonamnete urmeaza a expira in urmatoarele 30 de zile
culoare verde pentru ambele abonamente valabile
"""


def combina_taguri(tag_amef, tag_gprs):
    taguri = {tag_amef, tag_gprs}
    if "expirat" in taguri:
        return "expirat"
    if "avertizare" in taguri:
        return "avertizare"
    return "valid"


"""
Functie pentru a calcula cat timp mai este pana la expirare
"""


def calculeaza_tag_abonament(data_exp):
    if not data_exp:
        return "expirat"

    if isinstance(data_exp, str):

        try:
            data_exp = datetime.fromisoformat(data_exp).date()
        except ValueError:
            return "expirat"

    azi = date.today()
    zile = (data_exp - azi).days

    if zile < 0:
        return "expirat"
    elif zile <= 30:
        return "avertizare"
    else:
        return "valid"


def calculeaza_tag_abonament_gprs(data_exp):
    return calculeaza_tag_abonament(data_exp)


"""
Functie de combinare a abonamentelor pentru un singur pop-up
"""
def afiseaza_lista_abonamente(parent, rows, tip):
    rows_filtrate = rows.copy()

    azi = date.today()

    # =========================
    # CANVAS + SCROLL
    # =========================
    canvas = tk.Canvas(parent)
    scrollbar = tk.Scrollbar(parent, orient="vertical", command=canvas.yview)
    scroll_frame = tk.Frame(canvas)

    scroll_frame.bind(
        "<Configure>",
        lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
    )

    canvas.create_window((0, 0), window=scroll_frame, anchor="nw")
    canvas.configure(yscrollcommand=scrollbar.set)

    canvas.pack(side="left", fill="both", expand=True)
    scrollbar.pack(side="right", fill="y")

    # =========================
    # SEARCH BAR
    # =========================
    search_frame = tk.Frame(parent)
    search_frame.pack(fill="x", pady=5)

    tk.Label(search_frame, text="Filtrează după tehnician:", font=("Arial", 10)).pack(side="left", padx=5)



    # =========================
    # SCROLL MOUSE
    # =========================
    def on_mousewheel(event):
        canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")

    canvas.bind("<Enter>", lambda e: canvas.bind_all("<MouseWheel>", on_mousewheel))
    canvas.bind("<Leave>", lambda e: canvas.unbind_all("<MouseWheel>"))

    selected_label = {"widget": None, "bg": None}

    # =========================
    # AFISARE LISTA
    # =========================
    def afiseaza_rows(lista):
        for widget in scroll_frame.winfo_children():
            widget.destroy()

        for r in lista:
            data_exp = r["data_exp"]
            if not data_exp:
                continue

            if isinstance(data_exp, str):
                try:
                    data_exp = datetime.fromisoformat(data_exp).date()
                except ValueError:
                    continue

            zile_ramase = (data_exp - azi).days

            if zile_ramase < 0:
                text_status = "EXPIRAT"
                culoare = "#f28c8c"
            elif 0 <= zile_ramase <= 30:
                text_status = f"expiră în {zile_ramase} zile"
                culoare = "#fff3b0"
            else:
                continue

            descriere = "abonament service" if tip == "amef" else "comunicație GPRS"

            text = (
                f"{r['Nume_Firma']} (CUI: {r['Cui']}) | "
                f"Seria: {r['Serie_Amef']} | "
                f"{descriere} | {data_exp} → {text_status} | "
                f"Tehnician: {r.get('Tehnician','')}"
            )

            lbl = tk.Label(
                scroll_frame,
                text=text,
                bg=culoare,
                anchor="w",
                justify="left",
                font=("Arial", 10),
                pady=5
            )
            lbl.pack(fill="x", pady=2)

            # SELECTARE CLICK
            def on_click(event, label=lbl, bg=culoare, row=r):
                if selected_label["widget"]:
                    selected_label["widget"].configure(bg=selected_label["bg"])

                label.configure(bg="#9ecbff")
                selected_label["widget"] = label
                selected_label["bg"] = bg

                print("Selectat:", row)

            lbl.bind("<Button-1>", on_click)

    # =========================
    # FILTRARE
    # =========================
    def filtreaza(event=None):
        selected = combo_tehnician.get().lower()

        if selected == "":
            lista_filtrata = rows
        else:
            lista_filtrata = [
                r for r in rows
                if selected == (r.get("Tehnician", "") or "").lower()
            ]

        afiseaza_rows(lista_filtrata)

        nonlocal rows_filtrate
        rows_filtrate = lista_filtrata

    # afișare inițială
    afiseaza_rows(rows)

    search_var = tk.StringVar()
    combo_tehnician = ttk.Combobox(
        search_frame,
        textvariable=search_var,
        width=30,
        state="readonly"
    )

    combo_tehnician["values"] = [
        "",  # <- important: fără selecție = toate
        "POP CIPRIAN",
        "GRECU DAN"
    ]

    combo_tehnician.current(0)
    combo_tehnician.pack(side="left", padx=5)
    combo_tehnician.bind("<<ComboboxSelected>>", filtreaza)
    combo_tehnician.set("")

    # =========================
    # BUTON EXPORT PDF
    # =========================
    img = Image.open(resource_path("icons/pdf.png"))
    img = img.resize((20, 20))
    icon_pdf = ImageTk.PhotoImage(img)

    btn_export = tk.Button(
        parent,
        text="Exportă PDF",
        image=icon_pdf,
        compound="left",
        font=("Segoe UI", 10),
        bg="#f8f9fa",
        fg="#0d6efd",
        activebackground="#e9ecef",
        activeforeground="#212529",
        bd=1,
        relief="raised",
        highlightthickness=0,
        padx=10,
        pady=6,
        cursor="hand2",
        command=lambda: exporta_clienti_pdf(rows_filtrate, tip)  # IMPORTANT
    )
    btn_export.image = icon_pdf
    btn_export.pack(pady=5)

    def on_enter(e):
        btn_export["bg"] = "#e9f2ff"

    def on_leave(e):
        btn_export["bg"] = "#ffffff"

    btn_export.bind("<Enter>", on_enter)
    btn_export.bind("<Leave>", on_leave)


def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)


"""
Functie pentru a aparea in pop-ul cu abonamentele ce expira sau au expirat
"""


def alerta_abonamente_combinate():
    conn = conectare_db()
    cursor = conn.cursor()

    cursor.execute("""
        SELECT 
            d.Nume_Firma,
            d.Cui,
            d.Status_Firma,
            s.Status_Punct_Lucru,
            s.Serie_Amef,
            s.Data_Exp_Abon,
            s.Data_Exp_Gprs,
            s.Tehnician,
            s.Tip_Conect_Anaf,
            s.Status_AMEF,
            s.Observatii
        FROM tabela_date_clienti d
        LEFT JOIN tabela_sedii_secundare s ON d.Nr_Crt = s.Id_Client
    """)

    rows = cursor.fetchall()
    conn.close()

    # pregătim datele separat
    amef_rows = []
    gprs_rows = []

    for r in rows:
        # filtram doar clientii activi
        status_firma = (r["Status_Firma"] or "").strip().upper()
        status_punct = (r["Status_Punct_Lucru"] or "").strip().upper()

        firma_inactiva = status_firma in [
            "INCHIS",
            "SUSPENDAT",
            "INACTIV-RENUNTAT"
        ]

        punct_inactiv = status_punct in [
            "INCHIS",
            "SUSPENDAT",
            "INACTIV"
        ]

        if firma_inactiva or punct_inactiv:
            continue

        if r["Data_Exp_Abon"]:
            amef_rows.append({
                "Nume_Firma": r["Nume_Firma"],
                "Cui": r["Cui"],
                "Serie_Amef": r["Serie_Amef"],
                "data_exp": r["Data_Exp_Abon"],
                "tip_conect_anaf": r["Tip_Conect_Anaf"],
                "observatii": r["Observatii"],
                "Tehnician": r["Tehnician"]
            })
        if r["Data_Exp_Gprs"]:
            gprs_rows.append({
                "Nume_Firma": r["Nume_Firma"],
                "Cui": r["Cui"],
                "Serie_Amef": r["Serie_Amef"],
                "data_exp": r["Data_Exp_Gprs"],
                "Tehnician": r["Tehnician"]
            })

    # fereastra popup
    popup = tk.Toplevel()
    popup.title("Alerte Abonamente")
    popup.geometry("1400x700")
    popup.state("zoomed")

    # ---------------- AMEF (sus) ----------------
    frame_amef = tk.LabelFrame(
        popup,
        text="Abonamente AMEF / Service",
        font=("Arial", 11, "bold"),
        padx=5,
        pady=5
    )
    frame_amef.pack(fill="both", expand=True, padx=10, pady=5)

    afiseaza_lista_abonamente(frame_amef, amef_rows, "amef")

    # ---------------- GPRS (jos) ----------------
    frame_gprs = tk.LabelFrame(
        popup,
        text="Abonamente Comunicație GPRS",
        font=("Arial", 11, "bold"),
        padx=5,
        pady=5
    )
    frame_gprs.pack(fill="both", expand=True, padx=10, pady=5)

    afiseaza_lista_abonamente(frame_gprs, gprs_rows, "gprs")


# Functie de export in pdf a clientilor cu abonamente pt luna curenta

pdfmetrics.registerFont(TTFont('ArialUnicode', 'arial.ttf'))  # asigură-te că ai arial.ttf în folder sau calea completă


def exporta_clienti_pdf(rows, tip):
    # Filtrăm doar ce expiră sau a expirat
    azi = date.today()
    rows_filtrate = []
    for r in rows:
        data_exp = r.get("data_exp")
        if not data_exp:
            continue

        if isinstance(data_exp, str):
            try:
                data_exp = datetime.fromisoformat(data_exp).date()
            except:
                continue

        zile_ramase = (data_exp - azi).days
        if zile_ramase < 0 or zile_ramase <= 30:  # doar expirat sau expira in 30 zile
            rows_filtrate.append(r)

    if not rows_filtrate:
        messagebox.showinfo("Export PDF", "Nu există date de exportat în intervalul de 30 zile!")
        return

    from tkinter import filedialog  # sus în fișier (o singură dată)

    file_path = filedialog.asksaveasfilename(
        defaultextension=".pdf",
        filetypes=[("PDF files", "*.pdf")],
        title="Salvează raportul PDF",
        initialfile=f"export_abonamente_{tip}.pdf"
    )

    if not file_path:
        return  # utilizatorul a apăsat Cancel

    doc = SimpleDocTemplate(
        file_path,
        pagesize=A4,
        rightMargin=15, leftMargin=15, topMargin=20, bottomMargin=20
    )

    styles = getSampleStyleSheet()
    styleN = styles["Normal"]
    styleN.fontName = "ArialUnicode"
    styleN.fontSize = 8
    styleN.leading = 10

    header = ["Nr", "Nume Firmă", "CUI", "Seria AMEF", "Tehnician", "Tip Abonament", "Data Expirare", "Status"]
    data = [[Paragraph(col, styleN) for col in header]]

    for idx, r in enumerate(rows_filtrate, start=1):
        data_exp = r.get("data_exp")
        if isinstance(data_exp, str):
            data_exp = datetime.fromisoformat(data_exp).date()
        zile_ramase = (data_exp - azi).days
        text_status = "EXPIRAT" if zile_ramase < 0 else f"expiră în {zile_ramase} zile"
        tip_descriere = "abonament service" if tip == "amef" else "comunicație GPRS"

        row_data = [
            str(idx),
            r.get("Nume_Firma", ""),
            r.get("Cui", ""),
            r.get("Serie_Amef", ""),
            r.get("Tehnician", "N/A"),
            tip_descriere,
            str(data_exp),
            text_status
        ]
        data.append([Paragraph(str(cell), styleN) for cell in row_data])

    # Calcul lățimi coloane
    page_width, _ = A4
    total_margin = doc.leftMargin + doc.rightMargin
    max_width = page_width - total_margin
    col_widths = [25, 140, 60, 60, 100, 100, 60, 60]
    sum_widths = sum(col_widths)
    if sum_widths > max_width:
        scale = max_width / sum_widths
        col_widths = [w * scale for w in col_widths]

    table = Table(data, colWidths=col_widths, repeatRows=1)
    table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, -1), 'ArialUnicode'),
        ('FONTSIZE', (0, 0), (-1, -1), 8),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.whitesmoke]),
    ]))

    try:
        doc.build([table])
        messagebox.showinfo("Export PDF", f"PDF-ul a fost generat cu succes: {file_path}")
    except Exception as e:
        messagebox.showerror("Eroare Export PDF", f"A apărut o eroare la generarea PDF: {e}")


# =========================
# Functia care populeaza campurile de date din cele 2 coloane date client si sediu/amef
# =========================
def populare_campuri_treeview(event):
    selected = tree.focus()
    if not selected:
        return

    values = tree.item(selected, "values")

    mapping = {
        # date client
        "CUI Client": values[2],  # Cod Fiscal
        "Nume firmă": values[1],  # Nume Firma
        "Adresă sediu": values[3],  # Sediu Social
        "Numar Telefon": values[4],  # Nr Telefon
        "Adresa mail": values[5],  # Mail
        "Registrul Comertului": values[6],  # Reg Comert
        "Plătitor TVA": values[7],  # Tva
        "Administrator Firma": values[8],  # Administrator
        "Status Firma": values[9],  # Statusul Firmei Activ/Inchis

        # sediu secundar
        "Punct de lucru": values[10],  # Punct Lucru
        "Status puncti lucru": values[11], # Statusul punctului dde lucru
        "Model Amef": values[12],  # Model AMEF
        "Serie Amef": values[13],  # Serie AMEF
        "Nui Amef": values[14],  # NUI
        "Tehnician Service": values[15],  # Tehnician srv
        "Data conectare Anaf": values[16],  # Data Conectare Anaf
        "Data expirare abonament": values[17],  # Data Exp. Abonament
        "Valoare contract - RON": values[18],  # Val_Ctr
        "Tip Abonament": values[19],  # Tip Abonament
        "Data expirare Gprs": values[20],  # Data expirarii comunicatie GPRS
        "Tipul conectarii": values[21], # Conectare GPRS sau LAN
        "Status AMEF": values[22], # Status punct lucru - inchis, suspendat,
        "Observatii": values[23] # comentarii sau informatii despre data cand am sunat clientul sau altele
    }

    def populate_widget(widget, val, var=None):
        """Populează un widget indiferent de tip (Entry, Combobox, DateEntry)"""
        if isinstance(widget, ttk.Combobox):
            state_orig = widget.cget("state")
            widget.config(state="normal")
            widget.set(val)
            widget.config(state=state_orig)
        elif isinstance(widget, DateEntry):
            try:
                widget.set_date(val)
            except:
                widget.set_date("")
        elif var is not None:
            var.set(val)  # Folosește StringVar direct pentru Serie Amef / NUI
        else:
            widget.delete(0, tk.END)
            widget.insert(0, val)

    # Populează atât frame-ul client, cât și sediu
    for label, val in mapping.items():
        if label in entries_client:
            populate_widget(entries_client[label], val)
        if label in entries_sediu:
            var = None
            if label == "Serie Amef":
                var = entry_serie_amef_var
            elif label == "Nui Amef":
                var = entry_nui_var
            populate_widget(entries_sediu[label], val, var)


"""
Functie pentru a modifica tehnicianul de service
daca clientul trece la alt tehnician
"""


def modifica_tehnician():
    serie_amef = entry_serie_amef.get().strip()
    tehnician_nou = entry_tehnician.get().strip()

    if not serie_amef:
        messagebox.showwarning("Eroare", "Trebuie să introduci seria AMEF pentru a identifica punctul de lucru")
        return

    if not tehnician_nou:
        messagebox.showwarning("Eroare", "Trebuie să introduci numele tehnicianului")
        return

    # conectare la baza de date
    conn = conectare_db()
    cursor = conn.cursor()

    try:
        # verificam daca exista punctul de lucru cu seria AMEF introdusa
        cursor.execute("""
        SELECT Id_Client, Punct_Lucru FROM tabela_sedii_secundare WHERE Serie_Amef=%s
        """, (serie_amef,))
        result = cursor.fetchone()

        if not result:
            messagebox.showinfo("Info", "Nu există niciun punct de lucru cu această serie AMEF")
            return

        id_client, punct_lucru = result

        # actualizare doar a tehnicianului pentru punctul de lucru respectiv
        cursor.execute("""
        UPDATE tabela_sedii_secundare
        SET Tehnician=%s
        WHERE Id_Client=%s AND Serie_Amef=%s
        """, (tehnician_nou, id_client, serie_amef))

        conn.commit()
        messagebox.showinfo("Succes", f"Numele tehnicianului a fost modificat pentru seria AMEF {serie_amef}!")


    except Exception as e:
        messagebox.showerror("Eroare", f"Nu s-a putut modifica tehnicianul: {e}")
    finally:
        conn.close()


#############################################################################
"""
Functie pentru export baza date in format CSV
La exportare vor aparea 3 campuri de export pentru cele 3 tabele din baza de date
"""


def export_csv():
    conn = conectare_db()
    cursor = conn.cursor()

    # export tabela_date_clienti
    cursor.execute("SELECT * FROM tabela_date_clienti")
    clienti = cursor.fetchall()
    clienti_headers = [i[0] for i in cursor.description]

    # export tabela_sedii_secundare
    cursor.execute("SELECT * FROM tabela_sedii_secundare")
    sedii = cursor.fetchall()
    sedii_headers = [i[0] for i in cursor.description]

    # export tabela istoric_abonamente
    cursor.execute("SELECT * FROM istoric_abonamente")
    istoric = cursor.fetchall()
    istoric_headers = [i[0] for i in cursor.description]

    conn.close()

    # alegem folder și nume fișier
    file_path = filedialog.asksaveasfilename(defaultextension=".csv",
                                             filetypes=[("CSV files", "*.csv")],
                                             title="Export Bază de Date")
    if not file_path:
        return

    # salvăm tabela_date_clienti
    with open(file_path.replace(".csv", "_clienti.csv"), "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(clienti_headers)
        writer.writerows(clienti)

    # salvăm tabela_sedii_secundare
    with open(file_path.replace(".csv", "_sedii.csv"), "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(sedii_headers)
        writer.writerows(sedii)

    # salvam istoric_abonamente
    with open(file_path.replace("csv", "_istoric.csv"), "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(istoric_headers)
        writer.writerows(istoric)

    messagebox.showinfo("Succes",
                        f"Baza de date a fost exportată:\n{file_path}_clienti.csv,  {file_path}_sedii.csv, {file_path}_istoric.csv")


"""
Functie pentru import baza date in format CSV
La import vor aparea 3 campuri de importare pentru cele 3 tabele din baza de date
"""


def import_csv():
    # alegem fișierele CSV
    file_clienti = filedialog.askopenfilename(title="Selectează CSV tabela_date_clienti",
                                              filetypes=[("CSV files", "*.csv")])
    if not file_clienti:
        return

    file_sedii = filedialog.askopenfilename(title="Selectează CSV tabela_sedii_secundare",
                                            filetypes=[("CSV files", "*.csv")])
    if not file_sedii:
        return

    file_istoric = filedialog.askopenfilename(title="Selecteaza CSV istoric_abonamente",
                                              filetypes=[("CSV files", "*.csv")])
    if not file_istoric:
        return

    conn = conectare_db()
    cursor = conn.cursor()

    # import tabela_date_clienti
    with open(file_clienti, newline="", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for row in reader:
            # verificăm dacă clientul există după CUI sau Reg_Comert
            cursor.execute("SELECT Nr_Crt FROM tabela_date_clienti WHERE Cui=%s OR Reg_Comert=%s",
                           (row['Cui'], row.get('Reg_Comert')))
            result = cursor.fetchone()
            if result:
                # update
                id_client = result[0]
                placeholders = ", ".join(f"{k}=%s" for k in row.keys() if k != "Nr_Crt")
                values = [row[k] for k in row.keys() if k != "Nr_Crt"]
                values.append(id_client)
                cursor.execute(f"UPDATE tabela_date_clienti SET {placeholders} WHERE Nr_Crt=%s", values)
            else:
                # insert
                columns = ", ".join(row.keys())
                placeholders = ", ".join(["%s"] * len(row))
                values = list(row.values())
                cursor.execute(f"INSERT INTO tabela_date_clienti ({columns}) VALUES ({placeholders})", values)

    # import tabela_sedii_secundare
    with open(file_sedii, newline="", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for row in reader:
            # verificăm dacă punctul de lucru există după Id_Client și Serie_Amef
            cursor.execute("SELECT 1 FROM tabela_sedii_secundare WHERE Id_Client=%s AND Serie_Amef=%s",
                           (row['Id_Client'], row['Serie_Amef']))
            if cursor.fetchone():
                # update
                placeholders = ", ".join(f"{k}=%s" for k in row.keys() if k not in ["Id_Client", "Serie_Amef"])
                values = [row[k] for k in row.keys() if k not in ["Id_Client", "Serie_Amef"]]
                values.extend([row['Id_Client'], row['Serie_Amef']])
                cursor.execute(f"UPDATE tabela_sedii_secundare SET {placeholders} WHERE Id_Client=%s AND Serie_Amef=%s",
                               values)
            else:
                # insert
                columns = ", ".join(row.keys())
                placeholders = ", ".join(["%s"] * len(row))
                values = list(row.values())
                cursor.execute(f"INSERT INTO tabela_sedii_secundare ({columns}) VALUES ({placeholders})", values)

    # import tabela istoric abonamente
    with open(file_istoric, newline="", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for row in reader:
            cursor.execute("SELECT 1 FROM istoric_abonamente WHERE id_sediu=%s AND tip_abonament=%s AND data_start=%s",
                           (row["id_sediu"], row["tip_abonament"], row["data_start"]))
            if cursor.fetchone():
                # update
                placeholders = ", ".join(
                    f"{k}=%s" for k in row.keys() if k not in ["id_sediu", "tip_abonament", "data_start"])
                values = [row[k] for k in row.keys() if k not in ["id_sediu", "tip_abonament", "data_start"]]
                values.extend([row["id_sediu"], row["tip_abonament"], row["data_start"]])
                cursor.execute(
                    f"UPDATE istoric_abonamente SET {placeholders} WHERE id_sediu=%s AND tip_abonament=%s AND data_start=%s",
                    values)
            else:
                # insert
                columns = ", ".join(row.keys())
                placeholders = ", ".join(["%s"] * len(row))
                values = list(row.values())
                cursor.execute(f"INSERT INTO istoric_abonamente ({columns}) VALUES ({placeholders})", values)

    conn.commit()
    conn.close()
    messagebox.showinfo("Succes", "Baza de date a fost importată cu succes!")


"""
Zona functiilor pentru istoricul abonamentelor service si gprs
Istoricul abonamentelor se va prelungi numai la buton prelungire abonamente
"""


def salveaza_istoric_abonament(id_client, id_sediu, serie_amef, tip_abonament, data_start, data_expirare,
                               observatii=""):
    conn = conectare_db()
    cursor = conn.cursor()
    cursor.execute("""
        INSERT INTO istoric_abonamente
        (id_client, id_sediu, serie_amef, tip_abonament, data_start, data_expirare, observatii)
        VALUES (?, ?, ?, ?, ?, ?, ?)
    """, (id_client, id_sediu, serie_amef, tip_abonament, data_start, data_expirare, observatii))
    conn.commit()
    conn.close()


def actualizeaza_sediu_secundar(id_sediu, tip_abonament, data_expirare):
    conn = conectare_db()
    cursor = conn.cursor()
    if tip_abonament == "SERVICE":
        cursor.execute("UPDATE tabela_sedii_secundare SET data_expirare_abonament=? WHERE Id=?",
                       (data_expirare, id_sediu))
    elif tip_abonament == "GPRS":
        cursor.execute("UPDATE tabela_sedii_secundare SET data_expirare_gprs=? WHERE Id=?",
                       (data_expirare, id_sediu))
    conn.commit()
    conn.close()


# --- Popup pentru prelungire abonament cu deplasare la 3 luni
def popup_prelungire_abonament_trimestrial(id_client, id_sediu, serie_amef, data_exp_service):
    popup = tk.Toplevel()
    popup.title("Prelungire abonament cu deplasare")
    popup.geometry("420x380")
    popup.grab_set()

    tk.Label(popup, text=f"Client: {id_client}", font=("Arial", 10, "bold")).pack(anchor="w", padx=10)
    tk.Label(popup, text=f"Serie AMEF: {serie_amef}", font=("Arial", 10, "bold")).pack(anchor="w", padx=10)

    # TIpul abonamentului adica cu deplasare la 3 luni buton creat numai pt a memora in istoric data deplasarii si incasarii
    tk.Label(popup, text="Abonament deplasare trimestrial")
    tip_var = tk.StringVar(value="SERVICE")

    # --- Data start ---
    tk.Label(popup, text="Data start prelungire").pack(pady=(10, 0))
    cal = DateEntry(popup, date_pattern="yyyy-mm-dd")
    cal.pack()

    def seteaza_data_initiala(*args):
        if tip_var.get() == "SERVICE" and data_exp_service:
            cal.set_date(data_exp_service)

    tip_var.trace_add("write", seteaza_data_initiala)
    seteaza_data_initiala()

    # --- CONFIRMA ---
    def confirma():
        tip = tip_var.get()
        data_start = cal.get_date()
        data_exp_noua = adauga_trei_luni(data_start)

        conn = conectare_db()
        cursor = conn.cursor()

        # Salvează ISTORIC
        cursor.execute("""
                INSERT INTO istoric_abonamente
                (id_client, id_sediu, serie_amef, tip_abonament,
                 data_start, data_expirare, observatii)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            """, (
            id_client, id_sediu, serie_amef,
            tip, data_start, data_exp_noua,
            "Prelungire manuală"
        ))

        # update sediu secundar
        if tip == "SERVICE":
            cursor.execute("""
                    UPDATE tabela_sedii_secundare
                    SET Data_Exp_Abon=?
                    WHERE Id_Client=? AND Serie_Amef=?
                """, (data_exp_noua, id_client, serie_amef))
        else:
            return

        conn.commit()
        conn.close()

        messagebox.showinfo("Succes", f"{tip} prelungit până la {data_exp_noua}")
        popup.destroy()
        cauta_in_treeview()  # refresh tabel

    tk.Button(
        popup,
        text="Prelungește cu 3 luni",
        bg="#cfe2f3",
        font=("Arial", 10, "bold"),
        command=confirma
    ).pack(pady=15)

# Final popup prelungire 3 luni


# Functie pentru prelungire 3 luni de abonament pentru clientii cu deplasare
def adauga_trei_luni(data):
    """
    Primește un obiect datetime.date și returnează data cu 3 luni adăugat
    """
    from dateutil.relativedelta import relativedelta
    # dacă e string, îl convertim la date
    if isinstance(data, str):
        try:
            data = datetime.fromisoformat(data).date()
        except ValueError:
            return None  # data invalidă
    return data + relativedelta(months=3)

# Final functie prelungire 3 luni

# Inceput functie buton prelungire 3  luni
def buton_prelungire_3_luni():
    selected = tree.selection()
    if not selected:
        messagebox.showwarning("Atenție", "Selectează un rând din tabel !")
        return
    row = tree.item(selected[0], "values")

    id_client = row[0]
    serie_amef = row[13]
    data_exp_service = row[17]
    # luăm id_sediu real
    conn = conectare_db()
    cursor = conn.cursor()
    cursor.execute("""
            SELECT Nr_Crt
            FROM tabela_sedii_secundare
            WHERE Id_Client=? AND Serie_Amef=?
        """, (id_client, serie_amef))
    result = cursor.fetchone()
    conn.close()

    if not result:
        messagebox.showerror("Eroare", "Sediu secundar negăsit!")
        return

    id_sediu = result["Nr_Crt"]

    popup_prelungire_abonament_trimestrial(
        id_client,
        id_sediu,
        serie_amef,
        data_exp_service
    )


# final functie buton prelungire 3 luni


# --- Popup pentru prelungire abonament anual ---
def popup_prelungire_abonament(id_client, id_sediu, serie_amef, data_exp_service, data_exp_gprs):
    popup = tk.Toplevel()
    popup.title("Prelungire abonament anual")
    popup.geometry("420x380")
    popup.grab_set()

    tk.Label(popup, text=f"Client: {id_client}", font=("Arial", 10, "bold")).pack(anchor="w", padx=10)
    tk.Label(popup, text=f"Serie AMEF: {serie_amef}", font=("Arial", 10)).pack(anchor="w", padx=10)

    # --- Tip abonament ---
    tk.Label(popup, text="Tip abonament anual").pack(pady=(10, 0))
    tip_var = tk.StringVar(value="SERVICE")
    cmb = ttk.Combobox(
        popup,
        textvariable=tip_var,
        values=["SERVICE", "GPRS"],
        state="readonly",
        width=20
    )
    cmb.pack()

    # --- Data start ---
    tk.Label(popup, text="Data start prelungire").pack(pady=(10, 0))
    cal = DateEntry(popup, date_pattern="yyyy-mm-dd")
    cal.pack()

    def seteaza_data_initiala(*args):
        if tip_var.get() == "SERVICE" and data_exp_service:
            cal.set_date(data_exp_service)
        elif tip_var.get() == "GPRS" and data_exp_gprs:
            cal.set_date(data_exp_gprs)

    tip_var.trace_add("write", seteaza_data_initiala)
    seteaza_data_initiala()

    # --- CONFIRMA ---
    def confirma():
        tip = tip_var.get()
        data_start = cal.get_date()
        data_exp_noua = adauga_un_an(data_start)

        conn = conectare_db()
        cursor = conn.cursor()

        # Salvează ISTORIC
        cursor.execute("""
            INSERT INTO istoric_abonamente
            (id_client, id_sediu, serie_amef, tip_abonament,
             data_start, data_expirare, observatii)
            VALUES (?, ?, ?, ?, ?, ?, ?)
        """, (
            id_client, id_sediu, serie_amef,
            tip, data_start, data_exp_noua,
            "Prelungire manuală"
        ))

        # update sediu secundar
        if tip == "SERVICE":
            cursor.execute("""
                UPDATE tabela_sedii_secundare
                SET Data_Exp_Abon=?
                WHERE Id_Client=? AND Serie_Amef=?
            """, (data_exp_noua, id_client, serie_amef))
        else:
            cursor.execute("""
                UPDATE tabela_sedii_secundare
                SET Data_Exp_Gprs=?
                WHERE Id_Client=? AND Serie_Amef=?
            """, (data_exp_noua, id_client, serie_amef))

        conn.commit()
        conn.close()

        messagebox.showinfo("Succes", f"{tip} prelungit până la {data_exp_noua}")
        popup.destroy()
        cauta_in_treeview()  # refresh tabel

    tk.Button(
        popup,
        text="Prelungește abonament anual",
        bg="#cfe2f3",
        font=("Arial", 10, "bold"),
        command=confirma
    ).pack(pady=15)


"""
Functie click dublu pe un client din campul de conectare
Si de aici la dublu click se deschide popup-ul de prelungire al abonamentului de service sau gprs
"""


def la_double_click(event):
    selected = tree.selection()
    if not selected:
        return

    row = tree.item(selected[0], "values")

    id_client = row[0]  # Nr_Crt client
    serie_amef = row[13]  # Serie AMEF
    data_exp_abon = row[17]  # Data expirare service
    data_exp_gprs = row[20]  # Data expirare gprs

    # luăm id_sediu real
    conn = conectare_db()
    cursor = conn.cursor()
    cursor.execute("""
        SELECT Nr_Crt
        FROM tabela_sedii_secundare
        WHERE Id_Client=? AND Serie_Amef=?
    """, (id_client, serie_amef))
    result = cursor.fetchone()
    conn.close()

    if not result:
        messagebox.showerror("Eroare", "Sediu secundar negăsit!")
        return

    id_sediu = result["Nr_Crt"]

    popup_prelungire_abonament(
        id_client,
        id_sediu,
        serie_amef,
        data_exp_abon,
        data_exp_gprs
    )


# --- Functia apelata la buton ---
def buton_prelungire():
    selected = tree.selection()
    if not selected:
        messagebox.showwarning("Atenție", "Selectează un rând din tabel !")
        return

    la_double_click(None)


# Functie adaugare un an de abonament din data selectata in pupu-pul de prelungire
def adauga_un_an(data):
    """
    Primește un obiect datetime.date și returnează data cu 1 an adăugat
    """
    from dateutil.relativedelta import relativedelta
    # dacă e string, îl convertim la date
    if isinstance(data, str):
        try:
            data = datetime.fromisoformat(data).date()
        except ValueError:
            return None  # data invalidă
    return data + relativedelta(years=1)


# Functie pentru copierea datelor cu click dreapta
right_click_event = None  # Variabila globala pentru click dreapta


def copy_selection(mode="cell", event=None):
    """Copiaza in clipboard celula sau randul selectat"""
    selected_items = tree.selection()
    if not selected_items:
        return

    clipboard_text = ""
    if mode == "row":
        # Copiere randuri complete
        for item in selected_items:
            values = tree.item(item)["values"]
            clipboard_text += "\t".join(str(v) for v in values) + "\n"
    elif mode == "cell":
        # Copiere celula cu click dreapta
        if event is None:
            return
        col = tree.identify_column(event.x)
        col_index = int(col.replace("#", "")) - 1
        for item in selected_items:
            value = tree.item(item)["values"][col_index]
            clipboard_text += str(value) + "\n"

    root.clipboard_clear()
    root.clipboard_append(clipboard_text.strip())
    print(f"Copied:\n{clipboard_text.strip()}")


# Funtie pentru afisare si cautare live istoric abonamente
def popup_istoric_abonamente():
    popup = tk.Toplevel(root)
    popup.title("Istoric abonamente")
    popup.geometry("1000x700")
    popup.grab_set()

    # ---------------- CAUTARE ----------------
    frame_cautare = tk.Frame(popup)
    frame_cautare.pack(fill="x", padx=10, pady=5)

    tk.Label(frame_cautare, text="Caută (Client / Serie AMEF / NUI):").pack(side="left")

    search_var = tk.StringVar()
    entry_search = tk.Entry(frame_cautare, textvariable=search_var, width=40)
    entry_search.pack(side="left", padx=5)

    # ---------------- TABEL ----------------
    frame_table = tk.Frame(popup)
    frame_table.pack(fill="both", expand=True, padx=10, pady=5)

    columns = (
        "client",
        "serie_amef",
        "nui",
        "tip",
        "data_start",
        "data_exp",
        "observatii"
    )

    tree = ttk.Treeview(frame_table, columns=columns, show="headings", selectmode="extended")
    tree.pack(fill="both", expand=True)

    headings = {
        "client": "Client",
        "serie_amef": "Serie AMEF",
        "nui": "NUI",
        "tip": "Tip Abonament",
        "data_start": "Data Start",
        "data_exp": "Data Expirare",
        "observatii": "Observații"
    }

    for col, txt in headings.items():
        tree.heading(col, text=txt)
        tree.column(col, width=120, anchor="w")

    # Scrollbar
    scrollbar = ttk.Scrollbar(frame_table, orient="vertical", command=tree.yview)
    scrollbar.pack(side="right", fill="y")
    tree.configure(yscrollcommand=scrollbar.set)

    # ---------------- DATE ----------------
    def incarca_date():
        tree.delete(*tree.get_children())

        conn = conectare_db()
        cursor = conn.cursor()

        cursor.execute("""
            SELECT 
                d.Nume_Firma AS client,
                i.serie_amef,
                s.nui,
                i.tip_abonament,
                i.data_start,
                i.data_expirare,
                i.observatii
            FROM istoric_abonamente i
            JOIN tabela_date_clienti d
                ON d.Nr_Crt = i.id_client
            LEFT JOIN tabela_sedii_secundare s
                ON s.Id_Client = d.Nr_Crt
                AND s.Serie_Amef = i.serie_amef
            ORDER BY i.data_start DESC;
        """)
        rows = cursor.fetchall()
        conn.close()

        for row in rows:
            tree.insert("", "end", values=(
                row["client"],
                row["serie_amef"],
                row["nui"],
                row["tip_abonament"],
                row["data_start"],
                row["data_expirare"],
                row["observatii"]
            ))

    # ---------------- FILTRARE ----------------
    def filtreaza(*args):
        query = search_var.get().lower()
        tree.delete(*tree.get_children())

        conn = conectare_db()
        cursor = conn.cursor()

        cursor.execute("""
            SELECT 
                d.Nume_Firma AS client,
                i.serie_amef,
                s.nui,
                i.tip_abonament,
                i.data_start,
                i.data_expirare,
                i.observatii
            FROM istoric_abonamente i
            JOIN tabela_date_clienti d
                ON d.Nr_Crt = i.id_client
            LEFT JOIN tabela_sedii_secundare s 
                ON s.Id_Client = d.Nr_Crt
                AND s.Serie_Amef = i.serie_amef
            WHERE 
                LOWER(d.Nume_Firma) LIKE ? OR
                LOWER(i.serie_amef) LIKE ? OR
                LOWER(s.nui) LIKE ?

            ORDER BY i.data_start DESC;
        """, (f"%{query}%", f"%{query}%", f"%{query}%"))

        rows = cursor.fetchall()
        conn.close()

        for row in rows:
            tree.insert("", "end", values=(
                row["client"],
                row["serie_amef"],
                row["nui"],
                row["tip_abonament"],
                row["data_start"],
                row["data_expirare"],
                row["observatii"]
            ))

    search_var.trace_add("write", filtreaza)

    incarca_date()

    # ----------------Buton pentru stergere istoric abonament------------
    # Butonul de ștergere
    tk.Button(
        popup,
        text="Șterge selecția",
        bg="#f28c8c",
        font=("Arial", 10, "bold"),
        command=lambda: sterge_selectie_istoric(tree)
    ).pack(pady=5)

    # ---------------- Buton inchidere pop-up istoric  ----------------
    tk.Button(
        popup,
        text="Închide",
        bg="#f28c8e",
        font=("Arial", 10, "bold"),
        command=popup.destroy
    ).pack(pady=8)


# Functie pentru stergerea istoricului
def sterge_selectie_istoric(tree):
    selected_items = tree.selection()
    if not selected_items:
        messagebox.showwarning("Atenție", "Nu ai selectat niciun rând!")
        return

    if not messagebox.askyesno("Confirmare", f"Sigur vrei să ștergi {len(selected_items)} rânduri?"):
        return

    try:
        conn = conectare_db()
        cursor = conn.cursor()

        for item in selected_items:
            values = tree.item(item, "values")
            print("Stergem:", values)
            client = values[0]
            serie_amef = values[1]
            tip_abonament = values[3]
            data_start = values[4]

            # ștergere după combinația unică
            cursor.execute("""
                DELETE FROM istoric_abonamente
                WHERE tip_abonament = ? AND serie_amef=? AND data_start=?              
            """, (tip_abonament, serie_amef, data_start))

            tree.delete(item)  # șterge și din Treeview
        conn.commit()
        messagebox.showinfo("Succes", f"{len(selected_items)} rânduri au fost șterse!")

    except Exception as e:
        messagebox.showerror("Eroare", f"A apărut o eroare la ștergere:\n{e}")
    finally:
        conn.close()

'''
Zona functiilor pentru generarea  documentelor pdf
declaratie instalare, pv defisca;izare, fisa service, contract service, etc
'''

def convert_docx_to_pdf(docx_path, pdf_path):

    try:
        import os
        import win32com.client
        os.system("taskkill /f /im WINWORD.EXE >nul 2>&1")

        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False

        doc = word.Documents.Open(os.path.abspath(docx_path))
        doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
        doc.Close(False)

        word.Quit()

    except Exception as e:
        print("Eroare conversie PDF:", e)


def resource_path(relative_path):
    """Returnează calea corectă și în exe și în python"""
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")

    return os.path.join(base_path, relative_path)

def insereaza_semnatura(doc, semnatura_path):
    for p in doc.paragraphs:
        if "{Semnatura}" in p.text:
            p.text = p.text.replace("{Semnatura}", "")

            if semnatura_path and os.path.exists(resource_path(semnatura_path)):
                run = p.add_run()
                run.add_picture(resource_path(semnatura_path), width=Inches(2))

    # 🔥 Tabele (ASTA ÎȚI LIPSEȘTE)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if "{Semnatura}" in p.text:
                        p.text = p.text.replace("{Semnatura}", "")
                        if semnatura_path:
                            run = p.add_run()
                            run.add_picture(resource_path(semnatura_path), width=Inches(1.2))

def nume_fisier_valid(text):
    text = str(text)
    text = re.sub(r'[\\/*?:"<>|]', "", text)  # elimina caractere interzise
    text = text.replace(" ", "_")
    return text

def get_client_selectat():
    selected = tree.selection()
    item = tree.item(selected[0])
    valori = item["values"]
    id_client = valori[0]
    serie_amef = valori[1] # coloana cu serie amef

    if not selected:
        messagebox.showwarning("Selectează client", "Selectează un client.")
        return None

    item = tree.item(selected[0])
    valori = item["values"]

    return valori[0]  # ID client

def get_date_client(id_client):

    conn = conectare_db()
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()

    cursor.execute("""
        SELECT 
            Administrator,
            Nume_Firma,
            Sediu_Social,
            Cui,
            Nr_Telefon         
        FROM tabela_date_clienti
        WHERE Nr_Crt = ?
    """, (id_client,))

    row = cursor.fetchone()
    conn.close()

    if not row:
        return None

    return dict(row)

def get_amef_client(id_client, serie_amef):

    conn = conectare_db()
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()

    # normalizează serie_amef
    serie_amef = str(serie_amef).strip().upper()
    print("ID client:", id_client)
    print("Serie AMEF selectată:", repr(serie_amef))

    cursor.execute("""
        SELECT Punct_Lucru, Status_Punct_Lucru, Serie_Amef, Model_Amef, Nui, Tehnician
        FROM tabela_sedii_secundare
        WHERE Id_Client = ? AND UPPER(TRIM(Serie_Amef)) = ?
    """, (id_client, serie_amef))

    row = cursor.fetchone()
    conn.close()

    if not row:
        return None

    return dict(row)

def citeste_excel_modele():

    df = pd.read_excel(resource_path("mapping_model_amef.xlsx"))
    df.columns = df.columns.str.strip().str.upper()

    modele = {}

    for _, row in df.iterrows():
        model = str(row["MODEL_AMEF"]).strip().upper()
        modele[model] = {
            "AVIZ_DISTRIBUTIE": row.get("AVIZ_DISTRIBUTIE", ""),
            "DATA_AVIZ": row.get("DATA_AVIZ", "")
        }
    return modele

def citeste_excel_tehnicieni():
    df = pd.read_excel(resource_path("mapping_tehnician.xlsx"))
    df.columns = df.columns.str.strip().str.upper()

    tehnicieni = {}
    for _, row in df.iterrows():
        nume = str(row["NUME"]).strip()
        tehnicieni[nume] = {
            "SIGILIU": row.get("SIGILIU", ""),
            "LEGITIMATIE": row.get("LEGITIMATIE", ""),
            "SEMNATURA": row.get("SEMNATURA", "")
        }
    return tehnicieni

# Functie pentru generarea declaratiei de instalare in pdf si docx
def genereaza_declaratie():
    progress = None
    progress_win = None
    try:
        # --- Selectie rând din Treeview ---
        selected = tree.selection()
        if not selected:
            messagebox.showwarning("Selectează client/serie", "Selectează un rând.")
            return

        item = tree.item(selected[0])
        valori = item["values"]

        # Presupunem coloanele: 0=Nr_Crt client, 1=Serie_Amef, 2=Model, ...
        id_client = valori[0]
        serie_amef = str(valori[13]).strip() # coloana serie amef din treeview

        # Preluare date din DB
        client = get_date_client(id_client)
        amef = get_amef_client(id_client, serie_amef)

        if not client or not amef:
            messagebox.showerror("Eroare", "Nu s-au găsit date pentru client/serie selectata.")
            return
        nume_firma = nume_fisier_valid(client["Nume_Firma"])


        # 2️⃣ citire Excel model amef
        modele_excel = citeste_excel_modele()
        tehnicieni_excel = citeste_excel_tehnicieni()

        # Mapare amef
        model_db = amef["Model_Amef"].strip().upper()
        aviz = modele_excel.get(model_db, {}).get("AVIZ_DISTRIBUTIE", "")
        data_aviz = modele_excel.get(model_db, {}).get("DATA_AVIZ", "")
        if data_aviz:
            if isinstance(data_aviz, datetime):
                data_aviz = data_aviz.strftime("%d-%m-%Y")

            else:
                try:
                    data_aviz = datetime.strptime(str(data_aviz), "%m/%d/%Y").strftime("%d-%m-%Y")
                except:
                    pass

        # Mapare Tehnician
        nume_tehnician = amef.get("Tehnician", "").strip().upper()
        sigiliu_tehnician = ""
        legitimatie_tehnician = ""

        # normalizează cheile din Excel
        tehnicieni_excel_norm = {k.strip().upper(): v for k, v in tehnicieni_excel.items()}

        data_azi = datetime.today().strftime("%d.%m.%Y")

        # Data creata pentru campul data achizitionarii amef
        data_achizitionare = (datetime.today() - timedelta(days=3)).strftime("%d.%m.%Y")

        if nume_tehnician in tehnicieni_excel_norm:
            sigiliu_tehnician = tehnicieni_excel_norm[nume_tehnician]["SIGILIU"]
            legitimatie_tehnician = tehnicieni_excel_norm[nume_tehnician]["LEGITIMATIE"]
            semnatura_tehnician = tehnicieni_excel_norm[nume_tehnician].get("SEMNATURA", "")
        else:
            print(f"Tehnician {nume_tehnician} nu a fost găsit în Excel")

        # --- Dicționar pentru template ---
        date = {
            "{Administrator}": client["Administrator"],
            "{Nume_Firma}": client["Nume_Firma"],
            "{Sediu_Social}": client["Sediu_Social"],
            "{Cui}": client["Cui"],
            "{Punct_Lucru}": amef["Punct_Lucru"],
            "{Serie_Amef}": amef["Serie_Amef"],
            "{Model_Amef}": amef["Model_Amef"],
            "{Nui}": amef["Nui"],
            "{Aviz_Distributie}": aviz,
            "{Data_Aviz}": data_aviz,
            "{Sigiliu}": sigiliu_tehnician,
            "{Tehnician}": nume_tehnician,
            "{Legitimatie}": legitimatie_tehnician,
            "{Data}": data_azi,
            "{Data_Achizitionare}": data_achizitionare
        }

        def replace_text_in_paragraph(paragraph, data):

            full_text = "".join(run.text for run in paragraph.runs)

            replaced = False

            for key, value in data.items():
                if key in full_text:
                    full_text = full_text.replace(key, str(value))
                    replaced = True

            if replaced:
                for run in paragraph.runs:
                    run.text = ""

                paragraph.runs[0].text = full_text

        template_path = resource_path("template/declaratie_instalare.docx")

        if not os.path.exists(template_path):
            messagebox.showerror(
                "Eroare",
                f"Template lipsă:\n{template_path}"
            )
            return

        doc = Document(template_path)
        # Paragrafe
        for p in doc.paragraphs:
            replace_text_in_paragraph(p, date)

        # -------------------------
        # Tabele
        # -------------------------
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        replace_text_in_paragraph(p, date)

        semnatura_path = tehnicieni_excel_norm.get(nume_tehnician, {}).get("SEMNATURA", "")
        insereaza_semnatura(doc, semnatura_path)

        # --- Dialog pentru a alege calea de salvare ---
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Document Word", "*.docx")],
            initialfile=f"Declaratie_instalare_{nume_firma}_{id_client}_{serie_amef}.docx",
            title="Salvează declarația"
        )

        if not file_path:  # utilizatorul a apasat Cancel
            return

        progress_win = tk.Toplevel(root)
        progress_win.title("Se generează...")
        progress_win.geometry("300x80")
        progress_win.resizable(False, False)

        label = tk.Label(
            progress_win,
            text="Se generează declarația..."
        )
        label.pack(pady=5)

        progress = ttk.Progressbar(
            progress_win,
            mode="indeterminate"
        )

        progress.pack(pady=5, padx=10, fill="x")

        progress.start()
        progress_win.update()

        # --- Salvare docx ---
        doc.save(file_path)

        # --- Optional: convertire PDF dacă ai funcția convert ---
        pdf_path = file_path.replace(".docx", ".pdf")
        try:
            convert_docx_to_pdf(file_path, pdf_path)

            if os.path.exists(pdf_path):
                messagebox.showinfo(
                    "Succes",
                    "Declaratia a fost generata in PDF si Doc"
                )

                # deschide folderul si selecteaza PDF-ul
                subprocess.Popen(
                    f'explorer /select,"{os.path.normpath(pdf_path)}"'
                )

            else:
                messagebox.showwarning(
                    "PDF",
                    "DOCX creat, PDF nu"
                )

        except Exception as e:
            messagebox.showerror(
                "Eroare PDF",
                str(e)
            )

        except Exception as e:
            messagebox.showerror(
                "Eroare generală",
                str(e)
            )

    finally:
        if progress:
            progress.stop()
            progress.destroy()
        if progress_win:
            progress_win.destroy()

# Final functie declaratie instalare

"""
Functie pentru generare pv defiscalizare din template
"""
def genereaza_pv_defiscalizare():
    progress = None
    progress_win = None
    try:
        # --- Selectie rând din Treeview ---
        selected = tree.selection()
        if not selected:
            messagebox.showwarning("Selectează client/serie", "Selectează un rând.")
            return

        item = tree.item(selected[0])
        valori = item["values"]

        # Presupunem coloanele: 0=Nr_Crt client, 1=Serie_Amef, 2=Model, ...
        id_client = valori[0]
        serie_amef = str(valori[13]).strip() # coloana serie amef din treeview

        # Preluare date din DB
        client = get_date_client(id_client)
        amef = get_amef_client(id_client, serie_amef)
        if not client or not amef:
            messagebox.showerror("Eroare", "Nu s-au găsit date pentru client/serie selectata.")
            return
        nume_firma = nume_fisier_valid(client["Nume_Firma"])

        # 2️⃣ citire Excel model amef
        modele_excel = citeste_excel_modele()
        tehnicieni_excel = citeste_excel_tehnicieni()

        # Mapare amef
        model_db = amef["Model_Amef"].strip().upper()
        aviz = modele_excel.get(model_db, {}).get("AVIZ_DISTRIBUTIE", "")
        data_aviz = modele_excel.get(model_db, {}).get("DATA_AVIZ", "")
        if data_aviz:

            if isinstance(data_aviz, datetime):
                data_aviz = data_aviz.strftime("%d-%m-%Y")

            else:
                try:
                    data_aviz = datetime.strptime(str(data_aviz), "%m/%d/%Y").strftime("%d-%m-%Y")
                except:
                    pass

        # Mapare Tehnician
        nume_tehnician = amef.get("Tehnician", "").strip().upper()
        sigiliu_tehnician = ""
        legitimatie_tehnician = ""

        # normalizează cheile din Excel
        tehnicieni_excel_norm = {k.strip().upper(): v for k, v in tehnicieni_excel.items()}

        data_azi = datetime.today().strftime("%d.%m.%Y")

        if nume_tehnician in tehnicieni_excel_norm:
            sigiliu_tehnician = tehnicieni_excel_norm[nume_tehnician]["SIGILIU"]
            legitimatie_tehnician = tehnicieni_excel_norm[nume_tehnician]["LEGITIMATIE"]
            semnatura_tehnician = tehnicieni_excel_norm[nume_tehnician].get("SEMNATURA", "")
        else:
            print(f"Tehnician {nume_tehnician} nu a fost găsit în Excel")

        # --- Dicționar pentru template ---
        date = {
            "{Administrator}": client["Administrator"],
            "{Nume_Firma}": client["Nume_Firma"],
            "{Sediu_Social}": client["Sediu_Social"],
            "{Cui}": client["Cui"],
            "{Punct_Lucru}": amef["Punct_Lucru"],
            "{Serie_Amef}": amef["Serie_Amef"],
            "{Model_Amef}": amef["Model_Amef"],
            "{Nui}": amef["Nui"],
            "{Aviz_Distributie}": aviz,
            "{Data_Aviz}": data_aviz,
            "{Sigiliu}": sigiliu_tehnician,
            "{Tehnician}": nume_tehnician,
            "{Legitimatie}": legitimatie_tehnician,
            "{Data}": data_azi

        }

        def replace_text_in_paragraph(paragraph, data):

            full_text = "".join(run.text for run in paragraph.runs)

            replaced = False

            for key, value in data.items():
                if key in full_text:
                    full_text = full_text.replace(key, str(value))
                    replaced = True

            if replaced:
                for run in paragraph.runs:
                    run.text = ""

                paragraph.runs[0].text = full_text


        # --- Înlocuire placeholder în paragrafe și tabele ---
        doc = Document(resource_path("template/pv_defiscalizare.docx"))
        # Paragrafe
        for p in doc.paragraphs:
            replace_text_in_paragraph(p, date)

        # -------------------------
        # Tabele
        # -------------------------
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        replace_text_in_paragraph(p, date)

        semnatura_path = tehnicieni_excel_norm.get(nume_tehnician, {}).get("SEMNATURA", "")
        insereaza_semnatura(doc, semnatura_path)

        # --- Dialog pentru a alege calea de salvare ---
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Document Word", "*.docx")],
            initialfile=f"Pv_Defiscalizare_{nume_firma}_{id_client}_{serie_amef}.docx",
            title="Salvează PV-ul"
        )

        if not file_path:  # utilizatorul a apasat Cancel
            return
        progress_win = tk.Toplevel(root)
        progress_win.title("Se generează...")
        progress_win.geometry("300x80")
        progress_win.resizable(False, False)

        label = tk.Label(
            progress_win,
            text="Se generează PV defiscalizare..."
        )
        label.pack(pady=5)

        progress = ttk.Progressbar(
            progress_win,
            mode="indeterminate"
        )

        progress.pack(pady=5, padx=10, fill="x")

        progress.start()
        progress_win.update()

        # --- Salvare docx ---
        doc.save(file_path)

        # --- Optional: convertire PDF dacă ai funcția convert ---
        pdf_path = file_path.replace(".docx", ".pdf")
        try:
            convert_docx_to_pdf(file_path, pdf_path)
            if os.path.exists(pdf_path):
                messagebox.showinfo(
                    "Succes",
                    "PV defiscalizare generatat cu succes in PDF si Doc"
                )

                # deschide folderul si selecteaza PDF-ul
                subprocess.Popen(
                    f'explorer /select,"{os.path.normpath(pdf_path)}"'
                )

            else:
                messagebox.showwarning(
                    "PDF",
                    "DOCX creat, PDF nu"
                )

        # try:
        #     os.system("taskkill /f /im WINWORD.EXE >nul 2>&1")

        except Exception as e:
            messagebox.showerror(
                "Eroare PDF",
                str(e)
            )


        except Exception as e:
            messagebox.showerror(
                "Eroare generală",
                str(e)
            )

    finally:
        if progress:
            progress.stop()
            progress.destroy()
        if progress_win:
            progress_win.destroy()


# Final functie generare pv defiscalizare din template


# Inceput functie pentru generare document de predare acte amef catre client
def genereaza_pv_predare_acte():
    progress = None
    progress_win = None
    try:
        # --- Selectie rând din Treeview ---
        selected = tree.selection()
        if not selected:
            messagebox.showwarning("Selectează client/serie", "Selectează un rând.")
            return

        item = tree.item(selected[0])
        valori = item["values"]

        # Presupunem coloanele: 0=Nr_Crt client, 1=Serie_Amef, 2=Model, ...
        id_client = valori[0]
        serie_amef = str(valori[13]).strip() # coloana serie amef din treeview

        # Preluare date din DB
        client = get_date_client(id_client)
        amef = get_amef_client(id_client, serie_amef)

        if not client or not amef:
            messagebox.showerror("Eroare", "Nu s-au găsit date pentru client/serie selectata.")
            return
        nume_firma = nume_fisier_valid(client["Nume_Firma"])

        # Mapare amef
        model_db = amef["Model_Amef"].strip().upper()

        data_azi = datetime.today().strftime("%d.%m.%Y")

        tehnicieni_excel = citeste_excel_tehnicieni()

        nume_tehnician = amef.get("Tehnician", "").strip().upper()

        tehnicieni_excel_norm = {k.strip().upper(): v for k, v in tehnicieni_excel.items()}

        semnatura_tehnician = ""

        if nume_tehnician in tehnicieni_excel_norm:
            semnatura_tehnician = tehnicieni_excel_norm[nume_tehnician].get("SEMNATURA", "")
        else:
            print(f"Tehnician {nume_tehnician} nu a fost găsit în Excel")


        # --- Dicționar pentru template ---
        date = {
            "{Administrator}": client["Administrator"],
            "{Nume_Firma}": client["Nume_Firma"],
            "{Sediu_Social}": client["Sediu_Social"],
            "{Cui}": client["Cui"],
            "{Nr_Telefon}": client["Nr_Telefon"],
            "{Punct_Lucru}": amef["Punct_Lucru"],
            "{Serie_Amef}": amef["Serie_Amef"],
            "{Model_Amef}": amef["Model_Amef"],
            "{Nui}": amef["Nui"],
            "{Data}": data_azi,
        }

        def replace_text_in_paragraph(paragraph, data):

            full_text = "".join(run.text for run in paragraph.runs)

            replaced = False

            for key, value in data.items():
                if key in full_text:
                    full_text = full_text.replace(key, str(value))
                    replaced = True

            if replaced:
                for run in paragraph.runs:
                    run.text = ""

                paragraph.runs[0].text = full_text

        # --- Înlocuire placeholder în paragrafe și tabele ---
        doc = Document(resource_path("template/pv_docs_amef.docx"))
        # Paragrafe
        for p in doc.paragraphs:
            replace_text_in_paragraph(p, date)


        # -------------------------
        # Tabele
        # -------------------------
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        replace_text_in_paragraph(p, date)

        insereaza_semnatura(doc, semnatura_tehnician)


        # --- Dialog pentru a alege calea de salvare ---
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Document Word", "*.docx")],
            initialfile=f"Confirmare_CI_{nume_firma}_{id_client}_{serie_amef}.docx",
            title="Salvează confirmare predare CI"
        )

        if not file_path:  # utilizatorul a apasat Cancel
            return

        progress_win = tk.Toplevel(root)
        progress_win.title("Se generează...")
        progress_win.geometry("300x80")
        progress_win.resizable(False, False)

        label = tk.Label(
            progress_win,
            text="Se generează pv predare..."
        )
        label.pack(pady=5)

        progress = ttk.Progressbar(
            progress_win,
            mode="indeterminate"
        )

        progress.pack(pady=5, padx=10, fill="x")

        progress.start()
        progress_win.update()

        # --- Salvare docx ---
        doc.save(file_path)

        # --- Optional: convertire PDF dacă ai funcția convert ---
        pdf_path = file_path.replace(".docx", ".pdf")
        try:
            convert_docx_to_pdf(file_path, pdf_path)
            if os.path.exists(pdf_path):
                messagebox.showinfo(
                    "Succes",
                    "Confirmarea a fost generata in PDF si Doc"
                )

                # deschide folderul si selecteaza PDF-ul
                subprocess.Popen(
                    f'explorer /select,"{os.path.normpath(pdf_path)}"'
                )

            else:
                messagebox.showwarning(
                    "PDF",
                    "DOCX creat, PDF nu"
                )

        except Exception as e:
            messagebox.showerror(
                "Eroare PDF",
                str(e)
            )

        except Exception as e:
            messagebox.showerror(
                "Eroare generală",
                str(e)
            )

    finally:
        if progress:
            progress.stop()
            progress.destroy()
        if progress_win:
            progress_win.destroy()

# Final functie generere document predare acte amef catre client


"""
Functie pentru generare fisa reparatie
"""
def genereaza_fisa_reparatie():
    progress = None
    progress_win = None
    try:
        # --- Selectie rând din Treeview ---
        selected = tree.selection()
        if not selected:
            messagebox.showwarning("Selectează client/serie", "Selectează un rând.")
            return

        item = tree.item(selected[0])
        valori = item["values"]

        # Presupunem coloanele: 0=Nr_Crt client, 1=Serie_Amef, 2=Model, ...
        id_client = valori[0]
        serie_amef = str(valori[13]).strip() # coloana serie amef din treeview

        # Preluare date din DB
        client = get_date_client(id_client)
        amef = get_amef_client(id_client, serie_amef)

        if not client or not amef:
            messagebox.showerror("Eroare", "Nu s-au găsit date pentru client/serie selectata.")
            return
        nume_firma = nume_fisier_valid(client["Nume_Firma"])

        # Mapare amef
        model_db = amef["Model_Amef"].strip().upper()

        data_azi = datetime.today().strftime("%d.%m.%Y")

        tehnicieni_excel = citeste_excel_tehnicieni()

        nume_tehnician = amef.get("Tehnician", "").strip().upper()

        tehnicieni_excel_norm = {k.strip().upper(): v for k, v in tehnicieni_excel.items()}

        semnatura_tehnician = ""

        if nume_tehnician in tehnicieni_excel_norm:
            semnatura_tehnician = tehnicieni_excel_norm[nume_tehnician].get("SEMNATURA", "")
        else:
            print(f"Tehnician {nume_tehnician} nu a fost găsit în Excel")


        # --- Dicționar pentru template ---
        date = {
            "{Administrator}": client["Administrator"],
            "{Nume_Firma}": client["Nume_Firma"],
            "{Sediu_Social}": client["Sediu_Social"],
            "{Cui}": client["Cui"],
            "{Nr_Telefon}": client["Nr_Telefon"],
            "{Punct_Lucru}": amef["Punct_Lucru"],
            "{Serie_Amef}": amef["Serie_Amef"],
            "{Model_Amef}": amef["Model_Amef"],
            "{Nui}": amef["Nui"],
            "{Data}": data_azi,
        }

        def replace_text_in_paragraph(paragraph, data):

            full_text = "".join(run.text for run in paragraph.runs)

            replaced = False

            for key, value in data.items():
                if key in full_text:
                    full_text = full_text.replace(key, str(value))
                    replaced = True

            if replaced:
                for run in paragraph.runs:
                    run.text = ""

                paragraph.runs[0].text = full_text

        # --- Înlocuire placeholder în paragrafe și tabele ---
        doc = Document(resource_path("template/fisa_service.docx"))
        # Paragrafe
        for p in doc.paragraphs:
            replace_text_in_paragraph(p, date)


        # -------------------------
        # Tabele
        # -------------------------
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        replace_text_in_paragraph(p, date)

        insereaza_semnatura(doc, semnatura_tehnician)


        # --- Dialog pentru a alege calea de salvare ---
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Document Word", "*.docx")],
            initialfile=f"Fisa_Service_{nume_firma}_{id_client}_{serie_amef}.docx",
            title="Salvează fisa reparatie"
        )

        if not file_path:  # utilizatorul a apasat Cancel
            return

        progress_win = tk.Toplevel(root)
        progress_win.title("Se generează...")
        progress_win.geometry("300x80")
        progress_win.resizable(False, False)

        label = tk.Label(
            progress_win,
            text="Se generează declarația..."
        )
        label.pack(pady=5)

        progress = ttk.Progressbar(
            progress_win,
            mode="indeterminate"
        )

        progress.pack(pady=5, padx=10, fill="x")

        progress.start()
        progress_win.update()

        # --- Salvare docx ---
        doc.save(file_path)

        # --- Optional: convertire PDF dacă ai funcția convert ---
        pdf_path = file_path.replace(".docx", ".pdf")
        try:
            convert_docx_to_pdf(file_path, pdf_path)
            if os.path.exists(pdf_path):
                messagebox.showinfo(
                    "Succes",
                    "Fisa Service fost generata in PDF si Doc"
                )

                # deschide folderul si selecteaza PDF-ul
                subprocess.Popen(
                    f'explorer /select,"{os.path.normpath(pdf_path)}"'
                )

            else:
                messagebox.showwarning(
                    "PDF",
                    "DOCX creat, PDF nu"
                )

        except Exception as e:
            messagebox.showerror(
                "Eroare PDF",
                str(e)
            )

        except Exception as e:
            messagebox.showerror(
                "Eroare generală",
                str(e)
            )

    finally:
        if progress:
            progress.stop()
            progress.destroy()
        if progress_win:
            progress_win.destroy()

# Final functie generare fisa reparatie


"""
Functie pentru generarea dosarului de asistenta tehnica
Un fel de carte interventie digitala
"""
def genereaza_dosar_asistenta():
    progress = None
    progress_win = None
    try:
        # --- Selectie rând din Treeview ---
        selected = tree.selection()
        if not selected:
            messagebox.showwarning("Selectează client/serie", "Selectează un rând.")
            return

        item = tree.item(selected[0])
        valori = item["values"]

        # Presupunem coloanele: 0=Nr_Crt client, 1=Serie_Amef, 2=Model, ...
        id_client = valori[0]
        serie_amef = str(valori[13]).strip()  # coloana serie amef din treeview

        # Preluare date din DB
        client = get_date_client(id_client)
        amef = get_amef_client(id_client, serie_amef)

        if not client or not amef:
            messagebox.showerror("Eroare", "Nu s-au găsit date pentru client/serie selectata.")
            return
        nume_firma = nume_fisier_valid(client["Nume_Firma"])

        # 2️⃣ citire Excel model amef
        modele_excel = citeste_excel_modele()
        tehnicieni_excel = citeste_excel_tehnicieni()

        # Mapare amef
        model_db = amef["Model_Amef"].strip().upper()
        aviz = modele_excel.get(model_db, {}).get("AVIZ_DISTRIBUTIE", "")
        data_aviz = modele_excel.get(model_db, {}).get("DATA_AVIZ", "")
        if data_aviz:
            if isinstance(data_aviz, datetime):
                data_aviz = data_aviz.strftime("%d-%m-%Y")

            else:
                try:
                    data_aviz = datetime.strptime(str(data_aviz), "%m/%d/%Y").strftime("%d-%m-%Y")
                except:
                    pass

        # Mapare Tehnician
        nume_tehnician = amef.get("Tehnician", "").strip().upper()
        sigiliu_tehnician = ""
        legitimatie_tehnician = ""

        # normalizează cheile din Excel
        tehnicieni_excel_norm = {k.strip().upper(): v for k, v in tehnicieni_excel.items()}

        data_azi = datetime.today().strftime("%d.%m.%Y")

        # Data creata pentru campul data achizitionarii amef
        data_achizitionare = (datetime.today() - timedelta(days=3)).strftime("%d.%m.%Y")

        if nume_tehnician in tehnicieni_excel_norm:
            sigiliu_tehnician = tehnicieni_excel_norm[nume_tehnician]["SIGILIU"]
            legitimatie_tehnician = tehnicieni_excel_norm[nume_tehnician]["LEGITIMATIE"]
            semnatura_tehnician = tehnicieni_excel_norm[nume_tehnician].get("SEMNATURA", "")
        else:
            print(f"Tehnician {nume_tehnician} nu a fost găsit în Excel")

        # --- Dicționar pentru template ---
        date = {
            "{Administrator}": client["Administrator"],
            "{Nume_Firma}": client["Nume_Firma"],
            "{Sediu_Social}": client["Sediu_Social"],
            "{Cui}": client["Cui"],
            "{Punct_Lucru}": amef["Punct_Lucru"],
            "{Serie_Amef}": amef["Serie_Amef"],
            "{Model_Amef}": amef["Model_Amef"],
            "{Nui}": amef["Nui"],
            "{Aviz_Distributie}": aviz,
            "{Data_Aviz}": data_aviz,
            "{Sigiliu}": sigiliu_tehnician,
            "{Tehnician}": nume_tehnician,
            "{Legitimatie}": legitimatie_tehnician,
            "{Data}": data_azi,
            "{Data_Achizitionare}": data_achizitionare
        }

        def replace_text_in_paragraph(paragraph, data):

            full_text = "".join(run.text for run in paragraph.runs)

            replaced = False

            for key, value in data.items():
                if key in full_text:
                    full_text = full_text.replace(key, str(value))
                    replaced = True

            if replaced:
                for run in paragraph.runs:
                    run.text = ""

                paragraph.runs[0].text = full_text

        template_path = resource_path("template/dosar_asistenta.docx")

        if not os.path.exists(template_path):
            messagebox.showerror(
                "Eroare",
                f"Template lipsă:\n{template_path}"
            )
            return

        doc = Document(template_path)
        # Paragrafe
        for p in doc.paragraphs:
            replace_text_in_paragraph(p, date)

        # -------------------------
        # Tabele
        # -------------------------
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        replace_text_in_paragraph(p, date)

        semnatura_path = tehnicieni_excel_norm.get(nume_tehnician, {}).get("SEMNATURA", "")
        insereaza_semnatura(doc, semnatura_path)

        # --- Dialog pentru a alege calea de salvare ---
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Document Word", "*.docx")],
            initialfile=f"Dosar_asistenta_{nume_firma}_{id_client}_{serie_amef}.docx",
            title="Salvează dosar asistenta"
        )

        if not file_path:  # utilizatorul a apasat Cancel
            return

        progress_win = tk.Toplevel(root)
        progress_win.title("Se generează...")
        progress_win.geometry("300x80")
        progress_win.resizable(False, False)

        label = tk.Label(
            progress_win,
            text="Se generează declarația..."
        )
        label.pack(pady=5)

        progress = ttk.Progressbar(
            progress_win,
            mode="indeterminate"
        )

        progress.pack(pady=5, padx=10, fill="x")

        progress.start()
        progress_win.update()

        # --- Salvare docx ---
        doc.save(file_path)

        # --- Optional: convertire PDF dacă ai funcția convert ---
        pdf_path = file_path.replace(".docx", ".pdf")
        try:
            convert_docx_to_pdf(file_path, pdf_path)

            if os.path.exists(pdf_path):
                messagebox.showinfo(
                    "Succes",
                    "Dosarul de asistenta a fost generat in PDF si Doc"
                )

                # deschide folderul si selecteaza PDF-ul
                subprocess.Popen(
                    f'explorer /select,"{os.path.normpath(pdf_path)}"'
                )

            else:
                messagebox.showwarning(
                    "PDF",
                    "DOCX creat, PDF nu"
                )

        except Exception as e:
            messagebox.showerror(
                "Eroare PDF",
                str(e)
            )

        except Exception as e:
            messagebox.showerror(
                "Eroare generală",
                str(e)
            )

    finally:
        if progress:
            progress.stop()
            progress.destroy()
        if progress_win:
            progress_win.destroy()

"""
Finalul functie pentru dosar asistenta
"""

"""
Functie pentru legenda explicativa culori in functie de statusul abonamentelor
"""
def creeaza_legenda_status(parent):
    frame = tk.Frame(parent)

    def bulina(color, text):
        row = tk.Frame(frame)
        row.pack(anchor="w", pady=2)

        canvas = tk.Canvas(row, width=14, height=14, highlightthickness=0)
        canvas.pack(side="left")
        canvas.create_oval(2, 2, 12, 12, fill=color, outline="")

        tk.Label(row, text=text, font=("Segoe UI", 9)).pack(side="left", padx=6)

    bulina("#008000", "Abonament in termen")  # Verde
    bulina("#f1c40f", "Abonamentul expiră în urmatoarele 30 de zile") # Galben
    bulina("#e74c3c", "Abonament expirat")  # Rosu
    bulina("#95a5a6", "Firma Inchisa, Suspendata, AMEF Defiscalizat, Renuntat") # Gri


    return frame




# =========================
# User Interface setup
# =========================
root = tk.Tk()
root.title("Gestionare Client și Sediu")
root.geometry("1400x700")


window_width = 1400
window_height = 700
# dimensiunea ecranului
screen_width = root.winfo_screenwidth()
screen_height = root.winfo_screenheight()

# Ajustăm dimensiunea ferestrei dacă e mai mare decât ecranul
if window_width > screen_width:
    window_width = screen_width - 50  # micim cu 50px pentru margin
if window_height > screen_height:
    window_height = screen_height - 50
# Calculăm coordonatele pentru a centra
x = int((screen_width / 2) - (window_width / 2))
y = int((screen_height / 2) - (window_height / 2))

# Setăm geometria finală
root.geometry(f"{window_width}x{window_height}+{x}+{y}")

color_client = "#d0e1f9"
color_sediu = "#f9f1d0"



# -------------------------
# FRAME CLIENT (stânga)
# -------------------------
frame_client = tk.LabelFrame(root, text="Date Client", bg=color_client, padx=10, pady=10, font=("Arial", 12, "bold"))
frame_client.grid(row=0, column=0, sticky="nsew", padx=10, pady=5)

client_labels = ["CUI Client", "Nume firmă", "Adresă sediu", "Registrul Comertului",
                 "Plătitor TVA", "Administrator Firma", "Status Firma", "Numar Telefon", "Adresa mail"]
entries_client = {}

for i, label in enumerate(client_labels):
    tk.Label(frame_client, text=label, bg=color_client, font=("Arial", 10)).grid(row=i, column=0, sticky="w", padx=5,
                                                                              pady=2)
    # Labeluri pentru coloararea tva
    label_tva = tk.Label(frame_client,text="TVA: -", font=("Arial", 14, "bold"), bg=color_client)
    label_tva.grid(row=3, column=2, rowspan=2, sticky="w", padx=10)

    # e = tk.Entry(frame_client, width=40)
    # e.grid(row=i, column=1, sticky="w", padx=5, pady=2)
    # Dropdown pt TVA

    """
    # comentam sectoru asta de cod ca preia da sau nu din api anaf
    if label == "Plătitor TVA":
        e = ttk.Combobox(frame_client, values=["DA", "NU"], state="readonly", width=37)
        # e.set("NU")
    """

    # Dropdown Status Firma ---> setam elif daca decomentam iful de mai sus
    if label == "Status Firma":
        e = ttk.Combobox(
            frame_client,
            values=["ACTIV", "INCHIS", "SUSPENDAT", "INACTIV-RENUNTAT"],
            state="readonly",
            width=37
        )
        e.set("ACTIV")

    else:
        e = tk.Entry(frame_client, width=40)

    e.grid(row=i, column=1, sticky="w", padx=5, pady=2)
    entries_client[label] = e

(entry_cui, entry_nume, entry_adresa, entry_reg_comert,
 entry_tva, entry_administrator, entry_status_firma,
 entry_telefon, entry_mail) = [entries_client[label] for label in client_labels]


# -------------------------------
# Înregistrare validatecommand
# -------------------------------
vcmd_nui = (root.register(validare_nui), "%P")
vcmd_serie = (root.register(validare_serie_amef), "%P")

# variabile asociate entry-urilor pentru validare live
entry_serie_amef_var = tk.StringVar()
entry_nui_var = tk.StringVar()

# -------------------------
# FRAME PUNCT LUCRU/AMEF (dreapta)
# -------------------------
frame_sediu = tk.LabelFrame(root, text="Sediu Secundar / AMEF", bg=color_sediu, padx=10, pady=10,
                            font=("Arial", 12, "bold"))
frame_sediu.grid(row=0, column=1, sticky="nsew", padx=10, pady=5)

sediu_labels = ["Punct de lucru", "Status punct lucru", "Model Amef", "Serie Amef", "Nui Amef",
                "Data conectare Anaf", "Tehnician Service", "Data expirare abonament",
                "Valoare contract - RON", "Tip Abonament", "Data expirare Gprs", "Tipul conectarii", "Status AMEF", "Observatii"]

entries_sediu = {}

for i, label in enumerate(sediu_labels):
    tk.Label(
        frame_sediu,
        text=label,
        bg=color_sediu,
        font=("Arial", 10)
    ).grid(row=i, column=0, sticky="w", padx=5, pady=2)

    if label in ("Data conectare Anaf", "Data expirare abonament", "Data expirare Gprs"):
        e = DateEntry(
            frame_sediu,
            width=37,
            date_pattern="yyyy-mm-dd"  # compatibil MySQL
        )

    # DROPDOWN MODEL AMEF
    elif label == "Model Amef":
        e = ttk.Combobox(
            frame_sediu,
            values=[
                "DATECS DP25 MX",
                "DATECS DP25",
                "DATECS WP50 MX",
                "DATECS WP50",
                "DATECS DP05 MX",
                "DATECS DP05",
                "DATECS DP150",
                "DATECS FP700",
                "DATECS FP800",
                "DAISY EXPERT SX",
                "DAISY COMPACT M",
                "DAISY COMPACT S",
                "TREMOL M20 EXCEL MASTER",
                "TREMOL M20 - VALMED",
                "TREMOL M20 ADPOS M",
                "TREMOL ACTIVA MINI",
                "TREMOL FP17-T810",
                "TREMOL S-ACTIVA",
                "ACTIVA GALAXY",
                "SUCCES M7",
                "CUSTOM BIG PLUS",
                "PARTNER 200",
                "PARTNER 600",
                "ZIT B20",
                "ZIT B30"
            ],
            state="readonly",
            width=37
        )

        # DROPDOWN  Status Punct lucru
    elif label == "Status punct lucru":
        e = ttk.Combobox(
            frame_sediu,
            values=[
                "ACTIV",
                "INCHIS",
                "SUSPENDAT",
                "INACTIV-RENUNTAT"
            ],
            state="readonly",
            width=37
        )
        e.set("ACTIV")

    # DROPDOWN  Status AMEF
    elif label == "Status AMEF":
        e = ttk.Combobox(
            frame_sediu,
            values=[
                "ACTIV",
                "DEFISCALIZAT-AMEF LA CLIENT",
                "DEFISCALZAT-RADIAT-PRELUAT SD",
                "LA CLIENT-FARA DEFISCALIZARE"
            ],
            state="readonly",
            width=37
        )
        e.set("ACTIV")

    # DROPDOWN TIP CONECTARE ANAF GPRS SAU LAN
    elif label == "Tipul conectarii":
        e = ttk.Combobox(
            frame_sediu,
            values=[
                "GPRS",
                "LAN",
                "WIFI"
            ],
            state="readonly",
            width=37
        )
        e.set("GPRS")

    # DROPDOWN TIP ABONAMENT
    elif label == "Tip Abonament":

        e = ttk.Combobox(
            frame_sediu,
            values=[
                "ANUAL",
                "DEPLASARE-INTERN",
                "DEPLASARE-EXTERN"
            ],
            state="readonly",
            width=37
        )
        e.set("ANUAL")

    # DROPDOWN TEHNICIAN SERVICE
    elif label == "Tehnician Service":

        e = ttk.Combobox(
            frame_sediu,
            values=[
                "POP CIPRIAN",
                "GRECU DAN"
            ],
            state="readonly",
            width=37
        )
        e.set("POP CIPRIAN")

    # ENTRY pentru Serie AMEF și NUI cu validare live

    elif label == "Serie Amef":
        e = tk.Entry(
            frame_sediu,
            textvariable=entry_serie_amef_var,
            width=40,
            validate="key",
            validatecommand=vcmd_serie
        )

    elif label == "Nui Amef":
        e = tk.Entry(
            frame_sediu,
            textvariable=entry_nui_var,
            width=40,
            validate="key",
            validatecommand=vcmd_nui
        )

    else:
        e = tk.Entry(frame_sediu, width=40)

    e.grid(row=i, column=1, sticky="w", padx=5, pady=2)
    entries_sediu[label] = e

(entry_punct_lucru, entry_status_punct_lucru, entry_model_amef, entry_serie_amef, entry_nui,
 entry_conectare_anaf, entry_tehnician, entry_data_exp, entry_val_ctr,
 entry_tip_abonament, entry_data_exp_gprs, entry_tip_conect_anaf, entry_status_amef, entry_observatii) = [entries_sediu[label] for label in sediu_labels]

"""
Pentru populare automata in functie de tip client platitor tva sau nu 
cu deplasare sau anual
"""
entry_tip_abonament.bind("<<ComboboxSelected>>", actualizeaza_valoare_contract)
entry_tva.bind("<<ComboboxSelected>>", actualizeaza_valoare_contract)

# -------------------------
# FRAME BUTOANE
# -------------------------

"""
Clasa care creaza hover cu mesajele deasupra butoanelor cand trecem cu mouseul peste ele
"""
class ToolTip:
    def __init__(self, widget, text):
        self.widget = widget
        self.text = text
        self.tip_window = None
        widget.bind("<Enter>", self.show_tip)
        widget.bind("<Leave>", self.hide_tip)

    def show_tip(self, event=None):
        if self.tip_window or not self.text:
            return
        x = self.widget.winfo_rootx() + 20
        y = self.widget.winfo_rooty() + self.widget.winfo_height() + 5
        self.tip_window = tw = tk.Toplevel(self.widget)
        tw.wm_overrideredirect(True)
        tw.wm_geometry(f"+{x}+{y}")
        label = tk.Label(tw, text=self.text, justify="left",
                         background="#ffffe0", relief="solid", borderwidth=1,
                         font=("Arial", 9))
        label.pack(ipadx=4, ipady=2)

    def hide_tip(self, event=None):
        if self.tip_window:
            self.tip_window.destroy()
            self.tip_window = None


frame_butoane = tk.Frame(root)
frame_butoane.grid(row=1, column=0, columnspan=3, pady=10)

# Frame pentru legenda culori
frame_btn = tk.Frame(frame_butoane)
frame_btn.grid(row=0, column=0)


frame_legenda = tk.Frame(frame_butoane)
frame_legenda.grid(row=0, column=1, padx=20, sticky="n")
legenda = creeaza_legenda_status(frame_legenda)
legenda.pack(anchor="n")

btn_params = [
    ("Caută cu API", lambda: cauta_firma(), "#d4f0d0", "Caută firma folosind API-ul ANAF"),
    ("Salvează client", lambda: salveaza_client(), "#cfe2f3", "Salvează clientul în baza de date locală"),
    ("Prelungeste 1 AN", lambda: buton_prelungire(), "#cfe2f3", "Prelungește abonamentul curent cu 1 an"),
    ("Prelungeste 3 luni", lambda: buton_prelungire_3_luni(), "#cfe2f3", "Prelungește abonamentul curent cu 3 luni"),
    #("Verifică TVA (ANAF)", lambda: webbrowser.open_new("https://www.anaf.ro/RegistruTVA/"), "#0000FF", "Deschide site-ul ANAF pentru verificare TVA"),
    ("Afiseaza Abonam.", lambda: alerta_abonamente_combinate(), "#ffd966", "Afișează alerta cu expirarea abonamentelor"),
    ("Istoric Abonament", lambda: popup_istoric_abonamente(), "#008080", "Afișează istoricul abonamentelor Service și GPRS"),
    ("Resetare câmpuri", lambda: resetare_toate_campurile(), "#cfe2f3", "Resetează toate câmpurile din formular"),
    #("Merge DB (admin)", lambda: update_baza_protejat(), "#f4b183", "Combină și actualizează 2 baze de date"),
    ("Dosar Asistenta", lambda: genereaza_dosar_asistenta(), "#fce5cd", "Genereaza dosar asistenta tehnica in Pdf"),
    ("Genereaza DI", lambda: genereaza_declaratie(), "#d9ead3", "Generează declarație de instalare PDF"),
    ("Genereaza PV", lambda: genereaza_pv_defiscalizare(), "#d9ead3", "Generează PV defiscalizare in PDF"),
    ("Fisa Service", lambda: genereaza_fisa_reparatie(), "#d9ead3", "Generează fisa reparatie in PDF"),
    ("Confirmare CI", lambda: genereaza_pv_predare_acte(), "#d9ead3", "Confirmare primire documente AMEF"), # Modifica pentru formular carte interventie
    # ("Fișă intervenție", lambda: genereaza_document("fisa_interventie"), "#cfe2f3", "Generează fișă de service PDF"),
    # ("Contract service", lambda: genereaza_document("contract_service"), "#ead1dc", "Generează contract de service PDF"),
]
for i, (text, cmd, color, descriere) in enumerate(btn_params):
    btn = tk.Button(
        frame_btn,
        text=text,
        command=cmd,
        width=16,
        bg=color,
        font=("Arial", 10, "bold")
    )

    btn.grid(row=i // 4, column=i % 4, pady=5, padx=10)

    ToolTip(btn, descriere)


# -------------------------
# FRAME TREE + SEARCH (sub butoane)
# -------------------------
frame_tree = tk.Frame(root)
frame_tree.grid(row=3, column=0, columnspan=2, sticky="nsew", padx=10, pady=5)

# Mutăm câmpul de căutare aici și facem entry-ul mai mare
search_frame = tk.Frame(frame_tree)
search_frame.pack(fill="x", pady=5)

tk.Label(search_frame, text="Caută dupa Nume, Cui, Serie sau Nui:", font=("Arial", 10, "bold")).pack(side="left",
                                                                                                     padx=5)
search_entry = tk.Entry(search_frame, width=50)  # mai mare
search_entry.pack(side="left", padx=5)
tk.Button(search_frame, text="Caută", command=cauta_in_treeview, bg="#d4f0d0", width=12).pack(side="left", padx=5)
# tk.Button(search_frame, text="Resetează", command=incarca_dropdown_puncte, bg="#f0d0d0", width=12).pack(side="left", padx=5)

# Buton resetare camp cautare
tk.Button(search_frame, text="Resetează", command=resetare_camp_cautare, bg="#f0d0d0", width=12).pack(side="left",
                                                                                                      padx=5)

# =========================
# TREEVIEW REZULTATE (SUB CAUTARE)
# =========================

frame_tabel = tk.Frame(frame_tree)
frame_tabel.pack(fill="both", expand=True)

# search_frame.grid(row=3, column=0, columnspan=2, padx=10, pady=10, sticky="nsew")
scroll_y = tk.Scrollbar(frame_tabel, orient="vertical")
scroll_x = tk.Scrollbar(frame_tabel, orient="horizontal")

columns = (
    "Nr_Crt", "Nume Firma", "Cod Fiscal", "Sediu Social",
    "Nr Telefon", "Mail", "Reg Comert", "Tva", "Administrator", "Status Firma",
    "Punct Lucru", "Status Punct Lucru", "Model AMEF", "Serie AMEF", "NUI",
    "Tehnician srv", "Data Conectare Anaf", "Data Exp. Abonament", "Val_Ctr", "Tip Abonament", "Data Exp. Gprs",
    "Tip conectare", "Status AMEF", "Observatii"

)
tree = ttk.Treeview(
    frame_tabel,
    columns=columns,
    show="headings",
    yscrollcommand=scroll_y.set,
    xscrollcommand=scroll_x.set
)
tree.tag_configure("expirat", background="#f28c8c")  # roșu
tree.tag_configure("avertizare", background="#fff3b0")  # galben
tree.tag_configure("valid", background="#d4f7d4")  # verde
tree.tag_configure("status_inactiv", background="#808080")  # gri pentru firme inactive

scroll_y.config(command=tree.yview)
scroll_x.config(command=tree.xview)

scroll_y.pack(side="right", fill="y")
scroll_x.pack(side="bottom", fill="x")
tree.pack(fill="both", expand=True)
tree.bind("<<TreeviewSelect>>",
          populare_campuri_treeview)  # cu linia asta activam functia de populare campuri cand selectam din cautare
tree.bind("<Double-1>", la_double_click)
search_entry.bind("<KeyRelease>", lambda e: cauta_in_treeview())  # cautare live in treeview

for col in columns:
    tree.heading(col, text=col)
    tree.column(col, width=130, anchor="w")

# Meniu
meniu = tk.Menu(root)
root.config(menu=meniu)

# Meniul de import export a bazei de date cu cele 3 tabele in format csv
import_export_menu = tk.Menu(meniu, tearoff=0)
export_menu = tk.Menu(meniu, tearoff=0)
meniu.add_cascade(label="Importa/Exporta DB", background="lightblue", menu=import_export_menu)
import_export_menu.add_command(label="Importa baza date", background="lightblue", command=import_csv)
import_export_menu.add_command(label="Exporta baza date", background="lightblue", command=export_csv)

# Meniul de stergere client sau punct lucru
sterge_menu = tk.Menu(meniu, tearoff=0)
meniu.add_cascade(label="Sterge Client/Punct Lucru", background="lightblue", menu=sterge_menu)
sterge_menu.add_command(label="Sterge Client", background="lightblue", foreground="red", command=sterge_client)
sterge_menu.add_command(label="Sterge Punct Lucru", background="lightblue", foreground="red", command=sterge_punct)

# Meniul de combinare (merge) al 2 baze de date
merge_db = tk.Menu(meniu, tearoff=0)
meniu.add_cascade(label="Combina 2 baze de date", background="lightblue", menu=merge_db)
merge_db.add_command(label="Merge DB", background="lightblue", command=update_baza_protejat)

menu = tk.Menu(root, tearoff=0)
menu.add_command(label="Copiaza celula", command=lambda: copy_selection("cell", right_click_event))
menu.add_command(label="Copiaza tot randul", command=lambda: copy_selection("row", right_click_event))


def show_menu(event):
    global right_click_event
    right_click_event = event
    menu.tk_popup(event.x_root, event.y_root)


tree.bind("<Button-3>", show_menu)  # Button-3 = click dreapta pentru copiere

# CONFIGURARE GRID ROOT
root.grid_rowconfigure(3, weight=1)
root.grid_columnconfigure(0, weight=1)
root.grid_columnconfigure(1, weight=1)

# --- POP-UP ALERTĂ ABONAMENTE ---
# root.after(100, alerta_abonamente_color)  # rulează pop-up-ul automat după ce UI-ul principal e gata

footer = tk.Label(root, text="Designed by Pop Ciprian, © 2026 - Copywrite Edition",
                  font=("Arial", 8, "italic"), fg="gray")
footer.grid(row=4, column=0, columnspan=2, sticky="e", padx=10, pady=5)

root.mainloop()


